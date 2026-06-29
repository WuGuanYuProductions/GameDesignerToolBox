import streamlit as st
import pandas as pd
import graphviz
import plotly.express as px
import io
import copy
import re
import urllib.request
import os
import tempfile
import hashlib
import ssl
import json
import platform
import random
import datetime
from PIL import Image as PILImage, ImageDraw, ImageFont

try:
    from streamlit_quill import st_quill
    HAS_QUILL = True
except ImportError:
    HAS_QUILL = False

try:
    ssl._create_default_https_context = ssl._create_unverified_context
except Exception:
    pass

st.set_page_config(
    page_title="游戏策划工具箱",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
.editor-toolbar div[data-testid="column"] button {
    padding: 1px 2px !important;
    font-size: 11px !important;
    min-height: 24px !important;
    height: 24px !important;
    width: 100% !important;
    white-space: nowrap !important;
    overflow: hidden !important;
    text-overflow: ellipsis !important;
}
</style>
""", unsafe_allow_html=True)

@st.cache_data(show_spinner=False)
def get_local_image_path(url_or_path):
    if not url_or_path or pd.isna(url_or_path):
        return None
    url_str = str(url_or_path).strip()
    if not url_str:
        return None
    if url_str.startswith("http://") or url_str.startswith("https://"):
        try:
            ext = os.path.splitext(url_str)[1] or ".png"
            if "?" in ext:
                ext = ext.split("?")[0]
            if ext.lower() not in [".png", ".jpg", ".jpeg", ".gif"]:
                ext = ".png"
            hasher = hashlib.md5(url_str.encode('utf-8'))
            temp_dir = tempfile.gettempdir()
            local_path = os.path.join(temp_dir, f"ui_ref_{hasher.hexdigest()}{ext}")
            if not os.path.exists(local_path):
                urllib.request.urlretrieve(url_str, local_path)
            return local_path
        except Exception:
            return None
    elif os.path.exists(url_str):
        return url_str
    return None

@st.cache_data
def get_default_data():
    return [
        {"父系统": "主系统", "子系统/模块": "角色系统", "节点类型": "核心系统", "复杂度权重": 4, "功能描述": "负责角色属性、装备、技能等数据的展示与修改"},
        {"父系统": "角色系统", "子系统/模块": "装备模块", "节点类型": "子系统", "复杂度权重": 3, "功能描述": "处理装备的穿戴、强化与洗练"},
        {"父系统": "角色系统", "子系统/模块": "技能模块", "节点类型": "子系统", "复杂度权重": 3, "功能描述": "技能的解锁、升级与配置"},
        {"父系统": "主系统", "子系统/模块": "战斗系统", "节点类型": "核心系统", "复杂度权重": 5, "功能描述": "核心战斗逻辑、关卡数据结算"},
    ]

@st.cache_data
def get_default_sequence_data():
    return [
        {"步骤": 1, "玩家输入": "点击界面“十连抽卡”按钮", "功能输出": "前端校验代币是否充足，向服务端请求抽卡数据", "特殊表现": "播放启动音效，抽卡按钮进入置灰冷却状态"},
        {"步骤": 2, "玩家输入": "无（等待服务器返回数据）", "功能输出": "服务端计算随机权重，扣除代币并下发获得道具列表", "特殊表现": "界面播放魔法阵聚能光效，全屏震屏特效"},
        {"步骤": 3, "玩家输入": "点击屏幕任意位置跳过", "功能输出": "客户端提前渲染结算卡牌，更新背包数据并保存", "特殊表现": "卡牌卡背翻转，稀有金卡闪烁强烈粒子光芒，播放高光音效"},
        {"步骤": 4, "玩家输入": "点击“确定”返回主界面", "功能输出": "销毁抽卡3D场景 and 特效，恢复主界面状态", "特殊表现": "场景淡出转场，主页面代币数值平滑减少变化动画"}
    ]

@st.cache_data
def get_default_fsm_data():
    return [
        {"源状态": "空闲巡逻 (Patrol)", "源状态类型": "初始状态", "触发事件/条件": "发现玩家 (距离 < 15m)", "变迁动作 (Action)": "播放咆哮动画，锁定目标并向其奔跑", "目标状态": "警戒追击 (Chase)", "目标状态类型": "常规状态"},
        {"源状态": "警戒追击 (Chase)", "源状态类型": "常规状态", "触发事件/条件": "进入攻击范围 (距离 < 2m)", "变迁动作 (Action)": "重置普攻CD，抬手准备挥砍", "目标状态": "普通攻击 (Attack)", "目标状态类型": "常规状态"},
        {"源状态": "普通攻击 (Attack)", "源状态类型": "常规状态", "触发事件/条件": "攻击判定结束 且 玩家仍然存活", "变迁动作 (Action)": "恢复至追击步伐", "目标状态": "警戒追击 (Chase)", "目标状态类型": "常规状态"},
        {"源状态": "警戒追击 (Chase)", "源状态类型": "常规状态", "触发事件/条件": "玩家脱离仇恨范围 (距离 > 25m)", "变迁动作 (Action)": "脱战，进入无敌状态并快速回血", "目标状态": "回归出生点 (Return)", "目标状态类型": "常规状态"},
        {"源状态": "回归出生点 (Return)", "源状态类型": "常规状态", "触发事件/条件": "到达出生点", "变迁动作 (Action)": "清空仇恨，恢复巡逻步态", "目标状态": "空闲巡逻 (Patrol)", "目标状态类型": "初始状态"},
        {"源状态": "任何状态", "源状态类型": "常规状态", "触发事件/条件": "生命值归零 (HP <= 0)", "变迁动作 (Action)": "停止一切逻辑，播放死亡动画并掉落装备", "目标状态": "死亡消散 (Dead)", "目标状态类型": "结束状态"}
    ]

@st.cache_data
def get_default_ui_hierarchy_data():
    return [
        {"父级UI": "游戏主画布 (Root Canvas)", "当前UI/界面": "主城大厅 (Main Lobby)", "界面类型": "全屏主界面", "参考图URL": "https://images.unsplash.com/photo-1542751371-adc38448a05e?w=400&q=80", "核心组件/元素": "头像、等级、体力条、活动Banner、底部功能轮盘", "界面备注": "游戏主入口，后台常驻渲染，预加载 3D 角色展示底台"},
        {"父级UI": "主城大厅 (Main Lobby)", "当前UI/界面": "角色详情页 (Hero Detail)", "界面类型": "独立功能页", "参考图URL": "https://images.unsplash.com/photo-1511512578047-dfb367046420?w=400&q=80", "核心组件/元素": "角色3D卡片展示、属性面板、装备槽位、升级/强化页签", "界面备注": "全屏覆盖，打开时停止主城背景音乐，播放养成专属视听反馈"},
        {"父级UI": "角色详情页 (Hero Detail)", "当前UI/界面": "强化装备弹窗 (Enhance Pop)", "界面类型": "二级弹窗", "参考图URL": "https://images.unsplash.com/photo-1607604276583-eef5d076aa5f?w=400&q=80", "核心组件/元素": "材料消耗区、成功率百分比、强化按钮、属性变化对比", "界面备注": "模态弹出，带有半透明灰色遮罩，点击外部不可关闭"},
        {"父级UI": "角色详情页 (Hero Detail)", "当前UI/界面": "装备洗练提示 (Tooltip)", "界面类型": "浮窗/提示", "参考图URL": "https://images.unsplash.com/photo-1550745165-9bc0b252726f?w=400&q=80", "核心组件/元素": "属性洗练范围、极品属性概率预览、词条说明", "界面备注": "属于非阻塞UI层，点击界面任意空白区域即自动淡出销毁"},
        {"父级UI": "游戏主画布 (Root Canvas)", "当前UI/界面": "排位赛匹配页 (Match Panel)", "界面类型": "全屏主界面", "参考图URL": "https://images.unsplash.com/photo-1552820728-8b83bb6b773f?w=400&q=80", "核心组件/元素": "段位展示、匹配倒计时、队友列表、准备按钮", "界面备注": "通过主城大厅的‘开始战斗’功能跳转进入"}
    ]

@st.cache_data
def get_default_combat_data():
    return [
        {"父级招式": "动作根节点 (Action Root)", "当前招式/动作": "普通连招起手 (Light ATK 1)", "动作类型": "起手招式", "触发按键": "J (轻击)", "动画总帧数": 20, "伤害判定帧": "4-7", "战斗判定/属性": "打断, 削韧 10", "招式参考图URL": "https://images.unsplash.com/photo-1542751371-adc38448a05e?w=400&q=80", "设计备注": "起手普攻一段，判定范围为前方小扇形，为后续派生提供基础"},
        {"父级招式": "普通连招起手 (Light ATK 1)", "当前招式/动作": "普通二段 (Light ATK 2)", "动作类型": "派生连招", "触发按键": "J (轻击)", "动画总帧数": 28, "伤害判定帧": "8-12", "战斗判定/属性": "硬直, 削韧 15", "招式参考图URL": "https://images.unsplash.com/photo-1511512578047-dfb367046420?w=400&q=80", "设计备注": "快速平砍，判定范围为正前方120度扇形，可派生重击或闪避"},
        {"父级招式": "普通二段 (Light ATK 2)", "当前招式/动作": "普通三段 (Light ATK 3)", "动作类型": "派生连招", "触发按键": "J (轻击)", "动画总帧数": 35, "伤害判定帧": "12-16", "战斗判定/属性": "击退, 削韧 20", "招式参考图URL": "https://images.unsplash.com/photo-1542751371-adc38448a05e?w=400&q=80", "设计备注": "向前小突进平砍，收招硬直较长，命中后可使怪物后退一小步"},
        {"父级招式": "普通二段 (Light ATK 2)", "当前招式/动作": "连招派生：重劈 (Heavy Finisher)", "动作类型": "终结技", "触发按键": "K (重击)", "动画总帧数": 52, "伤害判定帧": "24-32", "战斗判定/属性": "霸体, 击飞, 削韧 50", "招式参考图URL": "https://images.unsplash.com/photo-1607604276583-eef5d076aa5f?w=400&q=80", "设计备注": "重击派生终结技，起手第10帧开始进入全身霸体，下砸造成大范围击飞"},
        {"父级招式": "普通三段 (Light ATK 3)", "当前招式/动作": "浮空极速刺击 (Air Thrust)", "动作类型": "蓄力/特殊", "触发按键": "J (蓄力1.5s)", "动画总帧数": 60, "伤害判定帧": "10-45", "战斗判定/属性": "无敌, 浮空追击, 削韧 5", "招式参考图URL": "https://images.unsplash.com/photo-1550745165-9bc0b252726f?w=400&q=80", "设计备注": "在空中对浮空状态的目标发动高速攒气穿刺，空中持续无敌状态"}
    ]

@st.cache_data
def get_default_fishbone_data():
    return [
        {"大类/主骨": "客户端 (Client)", "具体原因/小骨": "内存泄漏导致闪退", "严重程度": 5, "原因细述": "中端以下设备在游玩30分钟后极易崩溃"},
        {"大类/主骨": "客户端 (Client)", "具体原因/小骨": "发热严重与掉帧", "严重程度": 4, "原因细述": "未做分级画质渲染限制"},
        {"大类/主骨": "服务端 (Server)", "具体原因/小骨": "高并发排队雪崩", "严重程度": 5, "原因细述": "Redis缓存未做削峰导致连接暴增"},
        {"大类/主骨": "服务端 (Server)", "具体原因/小骨": "物理碰撞校验延迟", "严重程度": 4, "原因细述": "物理碰撞校验协议开销过大"},
        {"大类/主骨": "设计问题 (Design)", "具体原因/小骨": "新手期流失率极高", "严重程度": 5, "原因细述": "首场Boss战机制教学过于冗长枯燥"},
        {"大类/主骨": "设计问题 (Design)", "具体原因/小骨": "关卡体验卡点过死", "严重程度": 4, "原因细述": "强化石产出与消耗比例失衡"},
        {"大类/主骨": "运营宣传 (Operation)", "具体原因/小骨": "买量素材不匹配", "严重程度": 3, "原因细述": "实际游戏画面与广告相差过大"}
    ]

if "system_df" not in st.session_state:
    st.session_state.system_df = pd.DataFrame(get_default_data())
if "seq_df" not in st.session_state:
    st.session_state.seq_df = pd.DataFrame(get_default_sequence_data())
if "fsm_df" not in st.session_state:
    st.session_state.fsm_df = pd.DataFrame(get_default_fsm_data())
if "ui_df" not in st.session_state:
    st.session_state.ui_df = pd.DataFrame(get_default_ui_hierarchy_data())
if "combat_df" not in st.session_state:
    st.session_state.combat_df = pd.DataFrame(get_default_combat_data())
if "fishbone_df" not in st.session_state:
    st.session_state.fishbone_df = pd.DataFrame(get_default_fishbone_data())

tool_ids = ["system", "seq", "fsm", "ui", "combat", "fishbone"]
for _tid in tool_ids:
    if f"owner_{_tid}" not in st.session_state:
        st.session_state[f"owner_{_tid}"] = ""
    if f"scheme_title_{_tid}" not in st.session_state:
        st.session_state[f"scheme_title_{_tid}"] = ""
    if f"rte_{_tid}_summary_text" not in st.session_state:
        st.session_state[f"rte_{_tid}_summary_text"] = ""
    if f"rte_{_tid}_scheme_text" not in st.session_state:
        st.session_state[f"rte_{_tid}_scheme_text"] = ""
    if f"quill_{_tid}_summary" not in st.session_state:
        st.session_state[f"quill_{_tid}_summary"] = ""
    if f"quill_{_tid}_scheme" not in st.session_state:
        st.session_state[f"quill_{_tid}_scheme"] = ""
    if f"{_tid}_ver" not in st.session_state:
        st.session_state[f"{_tid}_ver"] = 0

if "export_ready" not in st.session_state:
    st.session_state.export_ready = False
if "export_bytes" not in st.session_state:
    st.session_state.export_bytes = None
if "export_filename" not in st.session_state:
    st.session_state.export_filename = ""
if "export_mime" not in st.session_state:
    st.session_state.export_mime = ""

def reset_export_state():
    st.session_state.export_ready = False
    st.session_state.export_bytes = None

def export_data_section(df, key, filename, sheet_name):
    export_format = st.radio("选择导出文件格式：", ["CSV 格式", "Excel (XLSX) 格式"], horizontal=True, key=f"fmt_{key}")
    if export_format == "CSV 格式":
        csv_data = df.to_csv(index=False).encode('utf-8-sig')
        st.download_button("📥 下载当前配置表 (.csv)", csv_data, f"{filename}.csv", "text/csv", width="stretch", key=f"dl_csv_{key}")
    else:
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name=sheet_name)
        st.download_button("📥 下载当前配置表 (.xlsx)", buffer.getvalue(), f"{filename}.xlsx", width="stretch", key=f"dl_xlsx_{key}")

@st.cache_data(show_spinner=False)
def cache_title_banner(text, level="H1"):
    return generate_title_banner_bytes(text, level)

@st.cache_data(show_spinner=False)
def cache_system_graph(df_json, rankdir):
    try:
        df = pd.read_json(io.StringIO(df_json))
        if df.empty:
            return None
        dot = graphviz.Digraph()
        dot.attr(dpi='400', rankdir=rankdir, bgcolor='#FFFFFF')
        rendered_nodes = set()
        preset_colors = [
            {"fill": "#E74C3C", "border": "#C0392B", "text": "white"},
            {"fill": "#3498DB", "border": "#2980B9", "text": "white"},
            {"fill": "#2ECC71", "border": "#27AE60", "text": "white"},
            {"fill": "#F1C40F", "border": "#F39C12", "text": "black"},
            {"fill": "#9B59B6", "border": "#8E44AD", "text": "white"}
        ]
        color_map = {n: preset_colors[i % len(preset_colors)] for i, n in enumerate(["核心系统", "子系统", "功能", "玩法", "其他"])}
        for _, row in df.iterrows():
            parent = str(row["父系统"]).strip()
            child = str(row["子系统/模块"]).strip()
            node_type = str(row["节点类型"]).strip()
            colors = color_map.get(node_type, {"fill": "#95A5A6", "border": "#7F8C8D", "text": "white"})
            if parent not in rendered_nodes:
                dot.node(parent, parent, shape='box', style='filled,rounded', fillcolor='#34495E', color='#2C3E50', fontcolor='white', fontname='Microsoft YaHei')
                rendered_nodes.add(parent)
            label = f"{child}\n({node_type})" if node_type else child
            dot.node(child, label, shape='box', style='filled,rounded', fillcolor=colors["fill"], color=colors["border"], fontcolor=colors["text"], fontname='Microsoft YaHei')
            rendered_nodes.add(child)
            dot.edge(parent, child, color='#7F8C8D', penwidth='1.5')
        return dot.pipe(format='png')
    except Exception:
        return None

@st.cache_data(show_spinner=False)
def cache_seq_graph(df_json, seq_direction):
    try:
        df = pd.read_json(io.StringIO(df_json))
        if df.empty:
            return None
        dot = graphviz.Digraph()
        dot.attr(dpi='400', rankdir=seq_direction, bgcolor='#FFFFFF')
        prev_id = None
        for idx, row in df.iterrows():
            step = row.get("步骤", idx + 1)
            player_in = str(row["玩家输入"]).strip() if pd.notna(row["玩家输入"]) else "无"
            sys_out = str(row["功能输出"]).strip() if pd.notna(row["功能输出"]) else "无"
            visual = str(row["特殊表现"]).strip() if pd.notna(row["特殊表现"]) else ""
            current_node_id = f"step_{idx}"
            label = f"""<
            <table border="0" cellborder="1" cellspacing="0" cellpadding="6">
            <tr><td bgcolor="#34495E"><font color="white"><b>步骤 {step}</b></font></td></tr>
            <tr><td bgcolor="#EBF5FB" align="left"><b>玩家操作:</b> {player_in}</td></tr>
            <tr><td bgcolor="#EAF2F8" align="left"><b>功能输出:</b> {sys_out}</td></tr>
            """
            if visual and visual != "无":
                label += f' <tr><td bgcolor="#FEF9E7" align="left"><b>表现效果:</b> {visual}</td></tr>'
            label += "</table>>"
            dot.node(current_node_id, label, shape='none', fontname='Microsoft YaHei')
            if prev_id is not None:
                dot.edge(prev_id, current_node_id, color='#3498DB', penwidth='2.0')
            prev_id = current_node_id
        return dot.pipe(format='png')
    except Exception:
        return None

@st.cache_data(show_spinner=False)
def cache_fsm_graph(df_json, fsm_dir):
    try:
        df = pd.read_json(io.StringIO(df_json))
        if df.empty:
            return None
        dot = graphviz.Digraph()
        dot.attr(dpi='400', rankdir=fsm_dir, bgcolor='#FFFFFF')
        state_colors = {
            "初始状态": {"fill": "#2ECC71", "border": "#27AE60", "text": "white"},
            "常规状态": {"fill": "#3498DB", "border": "#2980B9", "text": "white"},
            "结束状态": {"fill": "#E74C3C", "border": "#C0392B", "text": "white"}
        }
        defined_nodes = set()
        for _, row in df.iterrows():
            src = str(row["源状态"]).strip()
            src_type = str(row["源状态类型"]).strip()
            tgt = str(row["目标状态"]).strip()
            tgt_type = str(row["目标状态类型"]).strip()
            event = str(row["触发事件/条件"]).strip() if pd.notna(row["触发事件/条件"]) else ""
            action = str(row["变迁动作 (Action)"]).strip() if pd.notna(row["变迁动作 (Action)"]) else ""
            if src not in defined_nodes:
                col = state_colors.get(src_type, state_colors["常规状态"])
                dot.node(src, src, shape='ellipse', style='filled', fillcolor=col["fill"], color=col["border"], fontcolor=col["text"], fontname='Microsoft YaHei')
                defined_nodes.add(src)
            if tgt not in defined_nodes:
                col = state_colors.get(tgt_type, state_colors["常规状态"])
                dot.node(tgt, tgt, shape='ellipse', style='filled', fillcolor=col["fill"], color=col["border"], fontcolor=col["text"], fontname='Microsoft YaHei')
                defined_nodes.add(tgt)
            edge_lbl = f"{event}"
            if action and action != "无" and action != "nan":
                edge_lbl += f"\n(Do: {action})"
            dot.edge(src, tgt, label=edge_lbl, color='#7F8C8D', fontname='Microsoft YaHei', fontsize='10')
        return dot.pipe(format='png')
    except Exception:
        return None

@st.cache_data(show_spinner=False)
def cache_ui_graph(df_json, ui_rankdir):
    try:
        df = pd.read_json(io.StringIO(df_json))
        if df.empty:
            return None
        dot = graphviz.Digraph()
        dot.attr(dpi='400', rankdir=ui_rankdir, bgcolor='#FFFFFF')
        type_style_map = {
            "全屏主界面": {"style": "filled,rounded", "fill": "#A9DFBF", "color": "#27AE60", "shape": "box"},
            "独立功能页": {"style": "filled,rounded", "fill": "#AED6F1", "color": "#2980B9", "shape": "box"},
            "二级弹窗": {"style": "dashed,filled", "fill": "#F9E79F", "color": "#D35400", "shape": "box"},
            "浮窗/提示": {"style": "dotted,filled", "fill": "#FADBD8", "color": "#C0392B", "shape": "ellipse"}
        }
        node_configs = {}
        for _, row in df.iterrows():
            child_name = str(row["当前UI/界面"]).strip()
            ui_type = str(row["界面类型"]).strip()
            ref_img_path = str(row["参考图URL"]).strip() if pd.notna(row["参考图URL"]) else ""
            local_img_path = get_local_image_path(ref_img_path) if ref_img_path else None
            if local_img_path:
                local_img_path = os.path.abspath(local_img_path).replace("\\", "/")
                node_configs[child_name] = {"type": ui_type, "local_img": local_img_path, "is_defined": True}
        for _, row in df.iterrows():
            parent_name = str(row["父级UI"]).strip()
            if parent_name not in node_configs:
                node_configs[parent_name] = {"type": "容器层", "local_img": None, "is_defined": False}
        for name, config in node_configs.items():
            if not config["is_defined"]:
                dot.node(name, name, shape='box', style='filled,rounded', fillcolor='#34495E', color='#2C3E50', fontcolor='white', fontname='Microsoft YaHei')
            else:
                ui_type = config["type"]
                local_img = config["local_img"]
                node_style = type_style_map.get(ui_type, {"style": "filled", "fill": "#BDC3C7", "color": "#7F8C8D", "shape": "box"})
                if local_img:
                    html_label = f"""<
                    <table border="0" cellborder="1" cellspacing="0" cellpadding="4">
                    <tr><td bgcolor="{node_style['fill']}"><font color="black"><b>{name}</b></font></td></tr>
                    <tr><td bgcolor="#FFFFFF" fixedsize="true" width="160" height="120"><img scale="true" src="{local_img}"/></td></tr>
                    <tr><td bgcolor="#F2F3F4"><font point-size="9" color="#5D6D7E">{ui_type}</font></td></tr>
                    </table>>"""
                    dot.node(name, html_label, shape='none', style=node_style["style"], fontname='Microsoft YaHei')
                else:
                    lbl = f"{name}\n({ui_type})"
                    dot.node(name, lbl, shape=node_style["shape"], style=node_style["style"], fillcolor=node_style["fill"], color=node_style["color"], fontcolor='black', fontname='Microsoft YaHei')
        for _, row in df.iterrows():
            parent = str(row["父级UI"]).strip()
            child = str(row["当前UI/界面"]).strip()
            dot.edge(parent, child, color='#7F8C8D', penwidth='1.5')
        return dot.pipe(format='png')
    except Exception:
        return None

@st.cache_data(show_spinner=False)
def cache_combat_graph(df_json, combat_rankdir):
    try:
        df = pd.read_json(io.StringIO(df_json))
        if df.empty:
            return None
        dot = graphviz.Digraph()
        dot.attr(dpi='400', rankdir=combat_rankdir, bgcolor='#FFFFFF')
        combat_style_map = {
            "起手招式": {"style": "filled,rounded", "fill": "#A9DFBF", "color": "#27AE60"},
            "派生连招": {"style": "filled,rounded", "fill": "#AED6F1", "color": "#2980B9"},
            "终结技": {"style": "filled,rounded", "fill": "#FADBD8", "color": "#C0392B"},
            "蓄力/特殊": {"style": "dashed,filled", "fill": "#F9E79F", "color": "#D35400"},
            "闪避/防御": {"style": "dotted,filled", "fill": "#D5F5E3", "color": "#117A65"}
        }
        node_configs = {}
        for _, row in df.iterrows():
            child_name = str(row["当前招式/动作"]).strip()
            action_type = str(row["动作类型"]).strip()
            key_input = str(row["触发按键"]).strip() if pd.notna(row["触发按键"]) else "无"
            total_frames = str(row["动画总帧数"]).strip() if pd.notna(row["动画总帧数"]) else "-"
            hit_frames = str(row["伤害判定帧"]).strip() if pd.notna(row["伤害判定帧"]) else "-"
            props = str(row["战斗判定/属性"]).strip() if pd.notna(row["战斗判定/属性"]) else "-"
            ref_img_path = str(row["招式参考图URL"]).strip() if pd.notna(row["招式参考图URL"]) else ""
            local_img_path = get_local_image_path(ref_img_path) if ref_img_path else None
            if local_img_path:
                local_img_path = os.path.abspath(local_img_path).replace("\\", "/")
                node_configs[child_name] = {
                    "type": action_type, "key_input": key_input, "total_frames": total_frames,
                    "hit_frames": hit_frames, "props": props, "local_img": local_img_path, "is_defined": True
                }
        for _, row in df.iterrows():
            parent_name = str(row["父级招式"]).strip()
            if parent_name not in node_configs:
                node_configs[parent_name] = {
                    "type": "基础节点", "key_input": "N/A", "total_frames": "-",
                    "hit_frames": "-", "props": "-", "local_img": None, "is_defined": False
                }
        for name, config in node_configs.items():
            if not config["is_defined"]:
                dot.node(name, name, shape='box', style='filled,rounded', fillcolor='#34495E', color='#2C3E50', fontcolor='white', fontname='Microsoft YaHei')
            else:
                action_type = config["type"]
                key_input = config["key_input"]
                total_frames = config["total_frames"]
                hit_frames = config["hit_frames"]
                props = config["props"]
                local_img = config["local_img"]
                node_style = combat_style_map.get(action_type, {"style": "filled", "fill": "#BDC3C7", "color": "#7F8C8D"})
                if local_img:
                    html_label = f"""<
                    <table border="0" cellborder="1" cellspacing="0" cellpadding="4">
                    <tr><td bgcolor="{node_style['fill']}"><font color="black"><b>{name}</b></font></td></tr>
                    <tr><td bgcolor="#FFFFFF" fixedsize="true" width="160" height="120"><img scale="true" src="{local_img}"/></td></tr>
                    <tr><td bgcolor="#F2F3F4">
                    <font point-size="9" color="#2C3E50"><b>按键:</b> {key_input} | <b>帧:</b> {total_frames}f ({hit_frames})</font><br/>
                    <font point-size="9" color="#C0392B"><b>判定:</b> {props}</font>
                    </td></tr>
                    </table>>"""
                else:
                    html_label = f"""<
                    <table border="0" cellborder="1" cellspacing="0" cellpadding="4">
                    <tr><td bgcolor="{node_style['fill']}"><font color="black"><b>{name} ({action_type})</b></font></td></tr>
                    <tr><td bgcolor="#FFFFFF">
                    <font point-size="9" color="#2C3E50"><b>输入按键:</b> {key_input}</font><br/>
                    <font point-size="9" color="#2C3E50"><b>动画全长:</b> {total_frames} 帧 | <b>判定活跃帧:</b> {hit_frames}</font><br/>
                    <font point-size="9" color="#C0392B"><b>核心判定:</b> {props}</font>
                    </td></tr>
                    </table>>"""
                dot.node(name, html_label, shape='none', style=node_style["style"], fontname='Microsoft YaHei')
        for _, row in df.iterrows():
            parent = str(row["父级招式"]).strip()
            child = str(row["当前招式/动作"]).strip()
            dot.edge(parent, child, color='#7F8C8D', penwidth='1.5')
        return dot.pipe(format='png')
    except Exception:
        return None

@st.cache_data(show_spinner=False)
def cache_fishbone_graph(df_json, problem_title):
    try:
        df = pd.read_json(io.StringIO(df_json))
        if df.empty:
            return None
        dot = graphviz.Digraph()
        dot.attr(dpi='400', rankdir='LR', bgcolor='#FFFFFF')
        dot.node("head", problem_title, shape='box', style='filled,rounded', fillcolor='#E74C3C', color='#C0392B', fontcolor='white', fontname='Microsoft YaHei', penwidth='2.0')
        categories = df["大类/主骨"].unique()
        preset_colors = [
            {"fill": "#3498DB", "border": "#2980B9", "text": "white"},
            {"fill": "#2ECC71", "border": "#27AE60", "text": "white"},
            {"fill": "#F1C40F", "border": "#F39C12", "text": "black"},
            {"fill": "#9B59B6", "border": "#8E44AD", "text": "white"},
            {"fill": "#E67E22", "border": "#D35400", "text": "white"}
        ]
        cat_color_map = {cat: preset_colors[i % len(preset_colors)] for i, cat in enumerate(categories)}
        for cat in categories:
            colors = cat_color_map[cat]
            dot.node(cat, cat, shape='ellipse', style='filled', fillcolor=colors["fill"], color=colors["border"], fontcolor=colors["text"], fontname='Microsoft YaHei')
            dot.edge(cat, "head", color='#34495E', penwidth='2.5')
        for idx, row in df.iterrows():
            cat = str(row["大类/主骨"]).strip()
            sub_cause = str(row["具体原因/小骨"]).strip()
            desc = str(row["原因细述"]).strip() if pd.notna(row["原因细述"]) else ""
            node_id = f"sub_{idx}"
            label = f"{sub_cause}\n(Lv.{row.get('严重程度', 3)})"
            dot.node(node_id, label, shape='plaintext', fontname='Microsoft YaHei', fontsize='10', tooltip=desc)
            dot.edge(node_id, cat, color='#7F8C8D', penwidth='1.0', style='dashed')
        return dot.pipe(format='png')
    except Exception:
        return None

@st.cache_data(show_spinner=False)
def cache_sunburst_image(df_json):
    try:
        df = pd.read_json(io.StringIO(df_json))
        if df.empty:
            return None
        fig = px.sunburst(
            df,
            path=['父系统', '子系统/模块'],
            values='复杂度权重',
            color='节点类型',
            color_discrete_map={
                "核心系统": "#E74C3C", "子系统": "#3498DB",
                "功能": "#2ECC71", "玩法": "#F1C40F", "其他": "#9B59B6", "(?)": "#BDC3C7"
            }
        )
        fig.update_layout(
            paper_bgcolor='white',
            plot_bgcolor='white',
            margin=dict(t=10, l=10, r=10, b=10),
            width=600,
            height=600
        )
        return fig.to_image(format="png", scale=4)
    except Exception:
        return None

def compile_system_graph_bytes():
    return cache_system_graph(st.session_state.system_df.to_json(orient="records"), "LR")

def compile_seq_graph_bytes():
    return cache_seq_graph(st.session_state.seq_df.to_json(orient="records"), "TB")

def compile_fsm_graph_bytes():
    return cache_fsm_graph(st.session_state.fsm_df.to_json(orient="records"), "LR")

def compile_ui_graph_bytes():
    return cache_ui_graph(st.session_state.ui_df.to_json(orient="records"), "LR")

def compile_combat_graph_bytes():
    return cache_combat_graph(st.session_state.combat_df.to_json(orient="records"), "LR")

def compile_fishbone_graph_bytes():
    return cache_fishbone_graph(st.session_state.fishbone_df.to_json(orient="records"), st.session_state.get("scheme_title_fishbone", "").strip() or "未命名核心问题")

def compile_sunburst_image_bytes():
    return cache_sunburst_image(st.session_state.system_df.to_json(orient="records"))

def render_graph_actions(png_bytes, key, filename):
    if png_bytes:
        st.image(png_bytes, width="stretch")
        st.markdown("##### 📸 导出示意图")
        col1, col2 = st.columns(2)
        with col1:
            st.download_button("🖼️ 导出为 PNG", png_bytes, f"{filename}_full.png", "image/png", width="stretch", key=f"btn_p_{key}")
        with col2:
            st.download_button("🖼️ 导出为 JPG", png_bytes, f"{filename}.jpg", "image/jpeg", width="stretch", key=f"btn_j_{key}")
    else:
        st.info("💡 提示：拓扑图渲染失败，请检查数据配置。")

def render_bytes_graph_actions(png_preview, key, filename):
    render_graph_actions(png_preview, key, filename)

@st.dialog("💾 保存记录到本地")
def show_save_dialog():
    st.write("请勾选您需要备份到本地文件的工具记录：")
    save_sys = st.checkbox("🎮 系统层级架构拆解", value=True)
    save_seq = st.checkbox("⏱️ 功能时序流向拆解", value=False)
    save_fsm = st.checkbox("📋 有限状态机拆解", value=False)
    save_ui = st.checkbox("🎨 界面层级拆解", value=False)
    save_combat = st.checkbox("⚔️ 战斗系统连招拆解", value=False)
    save_fb = st.checkbox("🐟 鱼骨问题分析", value=False)
    st.markdown("---")
    selected_tags = []
    export_payload = {}
    if save_sys:
        selected_tags.append("🎮 系统层级架构拆解")
        export_payload["system_df"] = st.session_state.system_df.to_dict(orient="records")
    if save_seq:
        selected_tags.append("⏱️ 功能时序流向拆解")
        export_payload["seq_df"] = st.session_state.seq_df.to_dict(orient="records")
    if save_fsm:
        selected_tags.append("📋 有限状态机拆解")
        export_payload["fsm_df"] = st.session_state.fsm_df.to_dict(orient="records")
    if save_ui:
        selected_tags.append("🎨 界面层级拆解")
        export_payload["ui_df"] = st.session_state.ui_df.to_dict(orient="records")
    if save_combat:
        selected_tags.append("⚔️ 战斗系统连招拆解")
        export_payload["combat_df"] = st.session_state.combat_df.to_dict(orient="records")
    if save_fb:
        selected_tags.append("🐟 鱼骨问题分析")
        export_payload["fishbone_df"] = st.session_state.fishbone_df.to_dict(orient="records")
    
    for _tid in tool_ids:
        export_payload[f"owner_{_tid}"] = st.session_state.get(f"owner_{_tid}", "").strip()
        export_payload[f"scheme_title_{_tid}"] = st.session_state.get(f"scheme_title_{_tid}", "").strip()
        
        # 修复同步 bug：Quill 实际存储在带有版本后缀_{ver}的 key 里
        ver = st.session_state.get(f"{_tid}_ver", 0)
        
        summary_val = st.session_state.get(f"rte_{_tid}_summary_text", "")
        if not summary_val or str(summary_val).strip() == "":
            summary_val = st.session_state.get(f"quill_{_tid}_summary_{ver}", "")
        export_payload[f"summary_{_tid}"] = summary_val
        
        scheme_val = st.session_state.get(f"rte_{_tid}_scheme_text", "")
        if not scheme_val or str(scheme_val).strip() == "":
            scheme_val = st.session_state.get(f"quill_{_tid}_scheme_{ver}", "")
        export_payload[f"scheme_{_tid}"] = scheme_val

    if not selected_tags:
        st.warning("⚠️ 请至少选择一个工具模块以进行数据保存。")
    else:
        file_data = {
            "file_identifier": "game_designer_toolbox_save",
            "tags": selected_tags,
            "data": export_payload
        }
        json_bytes = json.dumps(file_data, ensure_ascii=False, indent=2).encode('utf-8')
        abbr_tags = [t[:2].strip(" 🎮⏱️📋🎨⚔️🐟") for t in selected_tags]
        filename = f"GameDesignDocumentsSave_{'_'.join(abbr_tags)}.json"
        st.download_button(
            label="💾 确认并下载备份 file",
            data=json_bytes,
            file_name=filename,
            mime="application/json",
            width="stretch"
        )

@st.dialog("📂 加载本地文档")
def show_load_dialog():
    st.write("### 导入设定")
    uploaded_file = st.file_uploader("请上传导出的 JSON 记录文件", type=["json"])
    if uploaded_file is not None:
        try:
            content = json.load(uploaded_file)
            if not isinstance(content, dict) or content.get("file_identifier") != "game_designer_toolbox_save":
                st.error("❌ 无法识别的文件结构！该文件并非由此工具箱导出的兼容备份，请重新选择。")
                return
            file_tags = content.get("tags", [])
            st.info(f"📋 检测到该备份文件内含的工具数据为：`{file_tags}`")
            
            selected_tools = st.multiselect(
                "选择要导入的工具记录（默认已全选文件内包含的工具）：",
                options=file_tags,
                default=file_tags
            )
            
            if not selected_tools:
                st.warning("⚠️ 请至少选择一个导入目标工具。")
                return
            
            if st.button("🔄 确认覆盖并导入数据", width="stretch"):
                data_block = content.get("data", {})
                
                if "🎮 系统层级架构拆解" in selected_tools and "system_df" in data_block:
                    st.session_state["system_df"] = pd.DataFrame(data_block["system_df"])
                    st.session_state["system_ver"] += 1
                    st.toast("🎮 系统层级架构数据已载入！", icon="✅")
                
                if "⏱️ 功能时序流向拆解" in selected_tools and "seq_df" in data_block:
                    st.session_state["seq_df"] = pd.DataFrame(data_block["seq_df"])
                    st.session_state["seq_ver"] += 1
                    st.toast("⏱️ 功能时序流向数据已载入！", icon="✅")
                
                if "📋 有限状态机拆解" in selected_tools and "fsm_df" in data_block:
                    st.session_state["fsm_df"] = pd.DataFrame(data_block["fsm_df"])
                    st.session_state["fsm_ver"] += 1
                    st.toast("📋 有限状态机数据已载入！", icon="✅")
                
                if "🎨 界面层级拆解" in selected_tools and "ui_df" in data_block:
                    st.session_state["ui_df"] = pd.DataFrame(data_block["ui_df"])
                    st.session_state["ui_ver"] += 1
                    st.toast("🎨 界面层级数据已载入！", icon="✅")
                
                if "⚔️ 战斗系统连招拆解" in selected_tools and "combat_df" in data_block:
                    st.session_state["combat_df"] = pd.DataFrame(data_block["combat_df"])
                    st.session_state["combat_ver"] += 1
                    st.toast("⚔️ 战斗系统连招数据已载入！", icon="✅")
                
                if "🐟 鱼骨问题分析" in selected_tools and "fishbone_df" in data_block:
                    st.session_state["fishbone_df"] = pd.DataFrame(data_block["fishbone_df"])
                    st.session_state["fishbone_ver"] += 1
                    st.toast("🐟 鱼骨问题分析数据已载入！", icon="✅")
                
                tool_name_to_id = {
                    "🎮 系统层级架构拆解": "system",
                    "⏱️ 功能时序流向拆解": "seq",
                    "📋 有限状态机拆解": "fsm",
                    "🎨 界面层级拆解": "ui",
                    "⚔️ 战斗系统连招拆解": "combat",
                    "🐟 鱼骨问题分析": "fishbone"
                }
                for tool_name in selected_tools:
                    _tid = tool_name_to_id.get(tool_name)
                    if _tid:
                        if f"owner_{_tid}" in data_block:
                            st.session_state[f"owner_{_tid}"] = data_block[f"owner_{_tid}"]
                        if f"scheme_title_{_tid}" in data_block:
                            st.session_state[f"scheme_title_{_tid}"] = data_block[f"scheme_title_{_tid}"]
                        
                        curr_ver = st.session_state.get(f"{_tid}_ver", 0)
                        next_ver = curr_ver + 1
                        st.session_state[f"{_tid}_ver"] = next_ver
                        
                        if f"summary_{_tid}" in data_block:
                            summary_val = data_block[f"summary_{_tid}"]
                            st.session_state[f"rte_{_tid}_summary_text"] = summary_val
                            st.session_state[f"quill_{_tid}_summary"] = summary_val
                            st.session_state[f"quill_{_tid}_summary_{next_ver}"] = summary_val
                            st.session_state[f"quill_just_imported_{_tid}_summary"] = True
                        if f"scheme_{_tid}" in data_block:
                            scheme_val = data_block[f"scheme_{_tid}"]
                            st.session_state[f"rte_{_tid}_scheme_text"] = scheme_val
                            st.session_state[f"quill_{_tid}_scheme"] = scheme_val
                            st.session_state[f"quill_{_tid}_scheme_{next_ver}"] = scheme_val
                            st.session_state[f"quill_just_imported_{_tid}_scheme"] = True
                
                st.rerun()
        except Exception as e:
            st.error(f"❌ 读取文件出错：{e}")

def clean_emoji(text):
    if not isinstance(text, str):
        return text
    emoji_pattern = re.compile(
        r'[\U00010000-\U0010ffff]|\u2600-\u27BF|[\u200d\uFE0F]',
        flags=re.UNICODE
    )
    return emoji_pattern.sub('', text)

def clean_df_emojis(df):
    cleaned_df = df.copy()
    for col in cleaned_df.columns:
        cleaned_df[col] = cleaned_df[col].apply(lambda x: clean_emoji(str(x)) if pd.notna(x) else "")
    return cleaned_df

def generate_title_banner_bytes(text, level="H1"):
    width = 1200 if level == "H1" else 800
    height = 80 if level == "H1" else 50
    img = PILImage.new("RGBA", (width, height), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    clean_text = clean_emoji(text).replace(":", "").replace("：", "").strip()
    theme_color = (46, 64, 83, 255)
    if "系统" in text or "关系" in text:
        theme_color = (231, 76, 60, 255)
    elif "时序" in text or "步骤" in text:
        theme_color = (52, 152, 219, 255)
    elif "状态" in text or "机器人" in text:
        theme_color = (46, 204, 113, 255)
    elif "界面" in text or "规划" in text:
        theme_color = (241, 196, 15, 255)
    elif "战斗" in text or "连招" in text:
        theme_color = (155, 89, 182, 255)
    elif "鱼骨" in text or "因果" in text or "问题" in text:
        theme_color = (142, 68, 173, 255)
    elif "数据" in text or "记录" in text:
        theme_color = (52, 73, 94, 255)
    elif "权重" in text or "占比" in text:
        theme_color = (211, 84, 0, 255)
    elif "图鉴" in text or "卡片" in text:
        theme_color = (22, 160, 133, 255)
    font_path = None
    current_os = platform.system()
    if current_os == "Windows":
        font_path = "C:/Windows/Fonts/msyhbd.ttc"
    elif current_os == "Darwin":
        font_path = "/System/Library/Fonts/PingFang.ttc"
    try:
        if font_path and os.path.exists(font_path):
            font = ImageFont.truetype(font_path, 34 if level == "H1" else 20)
        else:
            font = ImageFont.load_default()
    except Exception:
        font = ImageFont.load_default()
    if level == "H1":
        draw.rounded_rectangle([10, 15, 26, 65], radius=6, fill=theme_color)
        draw.text((45, 18), clean_text, fill=(44, 62, 80, 255), font=font)
        draw.line([10, 78, width - 10, 78], fill=(220, 220, 220, 255), width=2)
    else:
        draw.rounded_rectangle([10, 10, 20, 38], radius=4, fill=theme_color)
        draw.text((32, 12), clean_text, fill=(52, 73, 94, 255), font=font)
    byte_arr = io.BytesIO()
    img.save(byte_arr, format='PNG')
    return byte_arr.getvalue(), width, height

def normalize_html_styles(text):
    """
    清洗并规范化输入中的所有富文本HTML样式、Quill 行内样式和文本换行。
    完美规整为 100% 对应且对称闭合的 <font color="..." size="..."> 样式标签，
    以便 Word 与 PDF 生成时高保真提取样式，彻底规避因标签失衡导致 ReportLab 等解析崩溃的问题。
    """
    if not text:
        return ""
    
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("<strong>", "<b>").replace("</strong>", "</b>")
    text = text.replace("<em>", "<i>").replace("</em>", "</i>")
    text = text.replace("<s>", "<strike>").replace("</s>", "</strike>")
    
    parts = re.split(r'(<[^>]+>)', text)
    converted_tags_stack = []
    new_parts = []
    
    for part in parts:
        if not part:
            continue
        if part.startswith('<') and part.endswith('>'):
            tag_lower = part.lower()
            if (tag_lower.startswith('<span') or tag_lower.startswith('<font')) and not tag_lower.startswith('</'):
                tag_content = part[1:-1]
                color_val = None
                rgb_match = re.search(r'color:\s*rgb\((\d+),\s*(\d+),\s*(\d+)\)', tag_content)
                if rgb_match:
                    r, g, b = map(int, rgb_match.groups())
                    color_val = f"#{r:02x}{g:02x}{b:02x}"
                else:
                    hex_match = re.search(r'(?:color[:=]\s*["\']?)(#[0-9a-fA-F]{3,6})', tag_content)
                    if hex_match:
                        color_val = hex_match.group(1)
                    else:
                        ql_color_match = re.search(r'class="ql-color-([^"]+)"', tag_content)
                        if ql_color_match:
                            c_name = ql_color_match.group(1)
                            if re.match(r'^[0-9a-fA-F]{6}$', c_name):
                                color_val = f"#{c_name}"
                            elif c_name.startswith('#'):
                                color_val = c_name
                            else:
                                color_map = {
                                    "red": "#E74C3C", "orange": "#E67E22", "yellow": "#F1C40F",
                                    "green": "#2ECC71", "blue": "#3498DB", "purple": "#9B59B6"
                                }
                                color_val = color_map.get(c_name, c_name)
                
                size_val = None
                px_match = re.search(r'font-size:\s*([\d\.]+)px', tag_content)
                if px_match:
                    px_val = float(px_match.group(1))
                    size_val = round(px_val * 0.75)
                else:
                    pt_match = re.search(r'font-size:\s*([\d\.]+)pt', tag_content)
                    if pt_match:
                        size_val = int(float(pt_match.group(1)))
                    else:
                        ql_size_match = re.search(r'class="ql-size-([^"]+)"', tag_content)
                        if ql_size_match:
                            size_name = ql_size_match.group(1)
                            pt_map = {"small": 9, "large": 18, "huge": 24}
                            size_val = pt_map.get(size_name, 12)
                        else:
                            font_size_match = re.search(r'size=["\']?(\d+)["\']?', tag_content)
                            if font_size_match:
                                size_map = {"1": 9, "2": 10, "3": 12, "4": 14, "5": 18, "6": 24, "7": 32}
                                size_val = size_map.get(font_size_match.group(1), 12)
                
                attrs = []
                if color_val:
                    attrs.append(f'color="{color_val}"')
                if size_val:
                    attrs.append(f'size="{size_val}"')
                
                if attrs:
                    new_parts.append(f"<font {' '.join(attrs)}>")
                    converted_tags_stack.append("font")
                else:
                    new_parts.append("")
                    converted_tags_stack.append("empty")
            elif tag_lower == '</span>' or tag_lower == '</font>':
                if converted_tags_stack:
                    state = converted_tags_stack.pop()
                    if state == "font":
                        new_parts.append("</font>")
                    else:
                        new_parts.append("")
                else:
                    new_parts.append("")
            else:
                new_parts.append(part)
        else:
            new_parts.append(part)
            
    return "".join(new_parts)

def add_markdown_paragraph_to_word(doc, text):
    from docx.shared import Pt, RGBColor
    if not text or not text.strip():
        p = doc.add_paragraph()
        p.add_run("[ 暂无策划设计说明 ]").font.italic = True
        return
    
    # 彻底转换为平衡规整的标准标签
    text = normalize_html_styles(text)
    
    # 分割段落：把 P 标签和列表标签转化为换行符 [P_BREAK] 切分
    text = text.replace('<p>', '').replace('</p>', '[P_BREAK]')
    text = text.replace('<br>', '\n').replace('<br/>', '\n')
    text = text.replace('<li>', '• ').replace('</li>', '[P_BREAK]')
    text = re.sub(r'</?(ul|ol)>', '', text)
    
    paragraphs_raw = text.split('[P_BREAK]')
    
    for p_raw in paragraphs_raw:
        p_raw_strip = p_raw.strip()
        if not p_raw_strip and p_raw != "":
            doc.add_paragraph()  # 空白占位段落
            continue
        if not p_raw_strip:
            continue
        
        p = doc.add_paragraph()
        
        # 按照标签把普通文本与状态标签切分开
        parts = re.split(r'(<[^>]+>)', p_raw)
        
        bold = False
        italic = False
        underline = False
        strike = False
        
        # 引入完美样式栈（Stack）来支持多样式完美嵌套，解决未指定值被 None 覆盖的问题
        color_stack = []
        size_stack = []
        
        for part in parts:
            if not part:
                continue
            if part.startswith('<') and part.endswith('>'):
                tag = part.lower()
                if tag == '<b>': bold = True
                elif tag == '</b>': bold = False
                elif tag == '<i>': italic = True
                elif tag == '</i>': italic = False
                elif tag == '<u>': underline = True
                elif tag == '</u>': underline = False
                elif tag == '<strike>': strike = True
                elif tag == '</strike>': strike = False
                elif tag.startswith('<font'):
                    hex_match = re.search(r'color="([^"]+)"', part)
                    pt_match = re.search(r'size="([^"]+)"', part)
                    
                    c_val = hex_match.group(1) if hex_match else None
                    s_val = float(pt_match.group(1)) if pt_match else None
                    
                    color_stack.append(c_val)
                    size_stack.append(s_val)
                elif tag == '</font>':
                    if color_stack: color_stack.pop()
                    if size_stack: size_stack.pop()
            else:
                # 处理普通纯文字块内的传统 \n 换行
                sub_lines = part.split('\n')
                for idx, line in enumerate(sub_lines):
                    if idx > 0:
                        p.add_run().add_break()  # 换行显示
                    if not line:
                        continue
                    
                    # 识别 markdown 的行内轻量级标记
                    sub_parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|~~.*?~~)', line)
                    for sp in sub_parts:
                        if not sp:
                            continue
                        r_bold = bold
                        r_italic = italic
                        r_strike = strike
                        val = sp
                        if sp.startswith('**') and sp.endswith('**'):
                            r_bold = True
                            val = sp[2:-2]
                        elif sp.startswith('*') and sp.endswith('*'):
                            r_italic = True
                            val = sp[1:-1]
                        elif sp.startswith('~~') and sp.endswith('~~'):
                            r_strike = True
                            val = sp[2:-2]
                        
                        cleaned_val = re.sub(r'<[^>]+>', '', val)
                        r = p.add_run(cleaned_val)
                        r.bold = r_bold
                        r.font.italic = r_italic
                        r.font.underline = underline
                        r.font.strike = r_strike
                        
                        # 完美从样式栈顶部获取当前的嵌套效果
                        active_color = next((c for c in reversed(color_stack) if c is not None), None)
                        active_size = next((s for s in reversed(size_stack) if s is not None), None)
                        
                        # 渲染文字颜色效果
                        if active_color:
                            try:
                                h = active_color.lstrip('#')
                                r.font.color.rgb = RGBColor(*(int(h[i:i+2], 16) for i in (0, 2, 4)))
                            except Exception:
                                pass
                        # 渲染字号效果
                        if active_size:
                            try:
                                r.font.size = Pt(active_size)
                            except Exception:
                                pass

def md_to_pdf_html(text):
    if not text or not text.strip():
        return "<font color='#7F8C8D'><i>[ 暂无策划设计说明 ]</i></font>"
    text = clean_emoji(text)
    
    # 清洗规范化 HTML 样式，输出对称的 <font color="..." size="..."> 样式标签
    text = normalize_html_styles(text)
    
    # 块级标签转换
    text = text.replace('<p>', '').replace('</p>', '<br/>')
    text = text.replace('<br>', '<br/>').replace('<br/>', '<br/>')
    text = text.replace('<li>', '&bull; ').replace('</li>', '<br/>')
    text = re.sub(r'</?(ul|ol)>', '', text)
    
    # 将换行符换成 ReportLab 兼容的折行标签
    text = text.replace('\n', '<br/>')
    
    # 解析 markdown 轻量级格式
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
    text = re.sub(r'~~(.*?)~~', r'<strike>\1</strike>', text)
    return text

def register_chinese_font():
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    current_os = platform.system()
    possible_fonts = []
    if current_os == "Windows":
        possible_fonts = [
            ("MicrosoftYaHei", "C:/Windows/Fonts/msyh.ttc"),
            ("SimSun", "C:/Windows/Fonts/simsun.ttc"),
            ("SimHei", "C:/Windows/Fonts/simhei.ttf"),
            ("MicrosoftYaHei-Bold", "C:/Windows/Fonts/msyhbd.ttc")
        ]
    elif current_os == "Darwin":
        possible_fonts = [
            ("PingFang", "/System/Library/Fonts/PingFang.ttc"),
            ("STHeiti", "/System/Library/Fonts/STHeiti Light.ttc"),
            ("Songti", "/System/Library/Fonts/Supplemental/Songti.ttc")
        ]
    else:
        possible_fonts = [
            ("ZenHei", "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"),
            ("DroidSansFallback", "/usr/share/fonts/truetype/droid/DroidSansFallback.ttf"),
            ("NotoSansCJK", "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc")
        ]
    for font_name, path in possible_fonts:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(font_name, path))
                return font_name
            except Exception:
                continue
    try:
        from reportlab.pdfbase.cidfonts import CIDFont
        pdfmetrics.registerFont(CIDFont('STSong-Light'))
        return 'STSong-Light'
    except Exception:
        return 'Helvetica'

def create_report_word(selected_configs):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    doc = Document()
    def add_page_number(run):
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        r = run._r
        r.append(fldChar1)
        r.append(instrText)
        r.append(fldChar2)
        r.append(fldChar3)

    def add_bookmark(paragraph, name):
        p = paragraph._p
        b_id = str(random.randint(1000, 99999))
        start = OxmlElement('w:bookmarkStart')
        start.set(qn('w:id'), b_id)
        start.set(qn('w:name'), name)
        end = OxmlElement('w:bookmarkEnd')
        end.set(qn('w:id'), b_id)
        p.insert(0, start)
        p.append(end)

    def add_page_reference(paragraph, name):
        p = paragraph._p
        fldSimple = OxmlElement('w:fldSimple')
        fldSimple.set(qn('w:instr'), f'PAGEREF {name} \\h')
        r = OxmlElement('w:r')
        rText = OxmlElement('w:t')
        rText.text = "1"
        r.append(rText)
        fldSimple.append(r)
        p.append(fldSimple)

    tid_map = {
        "系统层级关系表": "system",
        "功能时序流向表": "seq",
        "有限状态机表": "fsm",
        "界面层级规划表": "ui",
        "战斗系统连招表": "combat",
        "鱼骨问题分析表": "fishbone"
    }
    p_toc_head = doc.add_paragraph()
    p_toc_head.alignment = 1
    p_toc_head.paragraph_format.space_before = Pt(18)
    p_toc_head.paragraph_format.space_after = Pt(18)
    run_toc_head = p_toc_head.add_run("目 录")
    run_toc_head.font.name = 'Microsoft YaHei'
    run_toc_head.font.size = Pt(20)
    run_toc_head.bold = True
    for item in selected_configs:
        raw_title = item["title"]
        parsed_title = raw_title[:-1] if raw_title.endswith("表") else raw_title
        current_tid = tid_map.get(raw_title, "system")
        scheme_title = st.session_state.get(f"scheme_title_{current_tid}", "").strip() or "游戏策划设计方案"
        p_m = doc.add_paragraph()
        p_m.paragraph_format.left_indent = Inches(0.4)
        p_m.paragraph_format.space_after = Pt(4)
        run_m = p_m.add_run(f"■ {scheme_title} ({parsed_title})")
        run_m.font.name = 'Microsoft YaHei'
        run_m.font.size = Pt(11.5)
        run_m.bold = True
        dots_run = p_m.add_run(" ............................................................................................................ ")
        dots_run.font.name = 'Microsoft YaHei'
        dots_run.font.size = Pt(9)
        dots_run.font.color.rgb = RGBColor(180, 180, 180)
        add_page_reference(p_m, f"bm_{current_tid}_main")
        
        sum_label = "问题总述" if current_tid == "fishbone" else "设计总述"
        sch_label = "迭代方案" if current_tid == "fishbone" else "设计方案"
        
        sub_sections = [
            (sum_label, f"bm_{current_tid}_summary"),
            (sch_label, f"bm_{current_tid}_scheme"),
            ("数据表单详细记录", f"bm_{current_tid}_data")
        ]
        if item["graph_fn"]():
            sub_sections.append(("实时架构与关系流向预览", f"bm_{current_tid}_graph"))
        for idx_sub, (sub_name, sub_bm) in enumerate(sub_sections):
            p_sub = doc.add_paragraph()
            p_sub.paragraph_format.left_indent = Inches(0.8)
            p_sub.paragraph_format.space_after = Pt(2)
            prefix = "├─ " if idx_sub < len(sub_sections) - 1 else "└─ "
            run_sub = p_sub.add_run(f"{prefix}{sub_name}")
            run_sub.font.name = 'Microsoft YaHei'
            run_sub.font.size = Pt(10)
            run_sub.font.color.rgb = RGBColor(127, 140, 141)
            dots_sub = p_sub.add_run(" ................................................................................................................. ")
            dots_sub.font.name = 'Microsoft YaHei'
            dots_sub.font.size = Pt(9)
            dots_sub.font.color.rgb = RGBColor(200, 200, 200)
            add_page_reference(p_sub, sub_bm)
    body_section = doc.add_section()
    body_section.footer.is_linked_to_previous = False
    sectPr = body_section._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:start'), '1')
    sectPr.append(pgNumType)
    footer = body_section.footer
    p_footer = footer.paragraphs[0]
    p_footer.alignment = 1
    p_footer.add_run("第 ").font.size = Pt(9)
    add_page_number(p_footer.add_run())
    p_footer.add_run(" 页").font.size = Pt(9)
    for idx, item in enumerate(selected_configs):
        if idx > 0:
            doc.add_page_break()
        raw_title = item["title"]
        parsed_title = raw_title[:-1] if raw_title.endswith("表") else raw_title
        current_tid = tid_map.get(raw_title, "system")
        scheme_title = st.session_state.get(f"scheme_title_{current_tid}", "").strip() or "游戏策划设计方案"
        p_main_title = doc.add_paragraph()
        p_main_title.alignment = 1
        p_main_title.paragraph_format.space_before = Pt(12)
        p_main_title.paragraph_format.space_after = Pt(4)
        run_main_title = p_main_title.add_run(scheme_title)
        run_main_title.font.name = 'Microsoft YaHei'
        run_main_title.font.size = Pt(22)
        run_main_title.bold = True
        add_bookmark(p_main_title, f"bm_{current_tid}_main")
        p_sub_title = doc.add_paragraph()
        p_sub_title.alignment = 1
        p_sub_title.paragraph_format.space_after = Pt(14)
        run_sub_title = p_sub_title.add_run(parsed_title)
        run_sub_title.font.name = 'Microsoft YaHei'
        run_sub_title.font.size = Pt(13)
        run_sub_title.font.color.rgb = RGBColor(127, 140, 141)
        
        owner_name = st.session_state.get(f"owner_{current_tid}", "").strip()
        p_owner = doc.add_paragraph()
        p_owner.alignment = 2
        p_owner.paragraph_format.space_after = Pt(4)
        run_owner = p_owner.add_run(f"负责人：{owner_name if owner_name else '未指定'}")
        run_owner.font.name = 'Microsoft YaHei'
        run_owner.font.size = Pt(11)
        run_owner.font.italic = True

        p_time = doc.add_paragraph()
        p_time.alignment = 2
        p_time.paragraph_format.space_after = Pt(18)
        run_time = p_time.add_run(f"迭代时间：{datetime.datetime.now().strftime('%Y/%m/%d')}")
        run_time.font.name = 'Microsoft YaHei'
        run_time.font.size = Pt(11)
        run_time.font.italic = True
        
        sum_label = "问题总述" if current_tid == "fishbone" else "设计总述"
        sch_label = "迭代方案" if current_tid == "fishbone" else "设计方案"

        p_h_summary = doc.add_paragraph()
        p_h_summary.paragraph_format.space_before = Pt(12)
        p_h_summary.paragraph_format.space_after = Pt(6)
        run_h_summary = p_h_summary.add_run(sum_label)
        run_h_summary.font.name = 'Microsoft YaHei'
        run_h_summary.font.size = Pt(14)
        run_h_summary.bold = True
        add_bookmark(p_h_summary, f"bm_{current_tid}_summary")
        
        ver = st.session_state.get(f"{current_tid}_ver", 0)
        summary_content = st.session_state.get(f"rte_{current_tid}_summary_text", "")
        if not summary_content or str(summary_content).strip() == "":
            summary_content = st.session_state.get(f"quill_{current_tid}_summary_{ver}", "")
        add_markdown_paragraph_to_word(doc, summary_content)
        
        p_h_scheme = doc.add_paragraph()
        p_h_scheme.paragraph_format.space_before = Pt(12)
        p_h_scheme.paragraph_format.space_after = Pt(6)
        run_h_scheme = p_h_scheme.add_run(sch_label)
        run_h_scheme.font.name = 'Microsoft YaHei'
        run_h_scheme.font.size = Pt(14)
        run_h_scheme.bold = True
        add_bookmark(p_h_scheme, f"bm_{current_tid}_scheme")
        scheme_content = st.session_state.get(f"rte_{current_tid}_scheme_text", "")
        if not scheme_content or str(scheme_content).strip() == "":
            scheme_content = st.session_state.get(f"quill_{current_tid}_scheme_{ver}", "")
        add_markdown_paragraph_to_word(doc, scheme_content)
        
        df = clean_df_emojis(item["df"])
        h2_data_banner, _, _ = cache_title_banner("📊 数据表单详细记录：", "H2")
        p_intro = doc.add_paragraph()
        p_intro.paragraph_format.space_before = Pt(18)
        p_intro.paragraph_format.space_after = Pt(8)
        p_intro.add_run().add_picture(io.BytesIO(h2_data_banner), width=Inches(2.5))
        add_bookmark(p_intro, f"bm_{current_tid}_data")
        table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        for col_idx, col_name in enumerate(df.columns):
            hdr_cells[col_idx].text = str(col_name)
            for paragraph in hdr_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
        for row_idx, row_vals in enumerate(df.itertuples(index=False), 1):
            row_cells = table.rows[row_idx].cells
            for col_idx, val in enumerate(row_vals):
                row_cells[col_idx].text = str(val)
                for paragraph in row_cells[col_idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9.5)
        graph_bytes = item["graph_fn"]()
        if graph_bytes:
            h2_graph_banner, _, _ = cache_title_banner("🎨 实时架构与关系流向预览：", "H2")
            p_img_title = doc.add_paragraph()
            p_img_title.paragraph_format.space_before = Pt(20)
            p_img_title.paragraph_format.space_after = Pt(8)
            p_img_title.paragraph_format.keep_with_next = True
            p_img_title.add_run().add_picture(io.BytesIO(h2_graph_banner), width=Inches(3.2))
            add_bookmark(p_img_title, f"bm_{current_tid}_graph")
            p_img_body = doc.add_paragraph()
            p_img_body.paragraph_format.space_after = Pt(10)
            p_img_body.add_run().add_picture(io.BytesIO(graph_bytes), width=Inches(6.0))
            if raw_title == "系统层级关系表":
                sun_bytes = compile_sunburst_image_bytes()
                if sun_bytes:
                    h2_sun_banner, _, _ = cache_title_banner("🎯 架构权重分析 (Sunburst 占比图)：", "H2")
                    p_sun_title = doc.add_paragraph()
                    p_sun_title.paragraph_format.space_before = Pt(20)
                    p_sun_title.paragraph_format.space_after = Pt(8)
                    p_sun_title.paragraph_format.keep_with_next = True
                    p_sun_title.add_run().add_picture(io.BytesIO(h2_sun_banner), width=Inches(3.6))
                    p_sun_body = doc.add_paragraph()
                    p_sun_body.add_run().add_picture(io.BytesIO(sun_bytes), width=Inches(4.2))
            elif raw_title == "界面层级规划表":
                valid_gallery = []
                for _, row_data in df.iterrows():
                    ref_url = row_data.get("参考图URL")
                    if pd.notna(ref_url) and str(ref_url).strip():
                        valid_gallery.append(row_data)
                if valid_gallery:
                    h2_gal_banner, _, _ = cache_title_banner("🖼️ 界面卡片详细图鉴：", "H2")
                    p_gal_title = doc.add_paragraph()
                    p_gal_title.paragraph_format.space_before = Pt(20)
                    p_gal_title.paragraph_format.space_after = Pt(8)
                    p_gal_title.paragraph_format.keep_with_next = True
                    p_gal_title.add_run().add_picture(io.BytesIO(h2_gal_banner), width=Inches(2.8))
                    for row_data in valid_gallery:
                        tbl = doc.add_table(rows=1, cols=2)
                        tbl.style = 'Table Grid'
                        tbl.columns[0].width = Inches(2.2)
                        tbl.columns[1].width = Inches(3.8)
                        cell_img = tbl.cell(0, 0)
                        img_path = get_local_image_path(str(row_data["参考图URL"]).strip())
                        if img_path and os.path.exists(img_path):
                            try:
                                cell_img.paragraphs[0].add_run().add_picture(img_path, width=Inches(2.0))
                            except Exception:
                                cell_img.paragraphs[0].text = "[图片解析失败]"
                        else:
                            cell_img.paragraphs[0].text = "[无可用图片]"
                        cell_txt = tbl.cell(0, 1)
                        cell_txt.paragraphs[0].add_run(f"🏷️ {row_data['当前UI/界面']}\n").bold = True
                        p_details = cell_txt.add_paragraph()
                        p_details.add_run(f"上级容器: {row_data['父级UI']}\n")
                        p_details.add_run(f"界面类型: {row_data['界面类型']}\n")
                        p_details.add_run(f"核心组件: {row_data.get('核心组件/元素', '无')}\n")
                        p_details.add_run(f"设计备注: {row_data.get('界面备注', '无')}")
                        doc.add_paragraph().paragraph_format.space_after = Pt(10)
            elif raw_title == "战斗系统连招表":
                valid_gallery = []
                for _, row_data in df.iterrows():
                    ref_url = row_data.get("招式参考图URL")
                    if pd.notna(ref_url) and str(ref_url).strip():
                        valid_gallery.append(row_data)
                if valid_gallery:
                    h2_gal_banner, _, _ = cache_title_banner("招式动作详细图鉴：", "H2")
                    p_gal_title = doc.add_paragraph()
                    p_gal_title.paragraph_format.space_before = Pt(20)
                    p_gal_title.paragraph_format.space_after = Pt(8)
                    p_gal_title.paragraph_format.keep_with_next = True
                    p_gal_title.add_run().add_picture(io.BytesIO(h2_gal_banner), width=Inches(2.8))
                    for row_data in valid_gallery:
                        tbl = doc.add_table(rows=1, cols=2)
                        tbl.style = 'Table Grid'
                        tbl.columns[0].width = Inches(2.2)
                        tbl.columns[1].width = Inches(3.8)
                        cell_img = tbl.cell(0, 0)
                        img_path = get_local_image_path(str(row_data["招式参考图URL"]).strip())
                        if img_path and os.path.exists(img_path):
                            try:
                                cell_img.paragraphs[0].add_run().add_picture(img_path, width=Inches(2.0))
                            except Exception:
                                cell_img.paragraphs[0].text = "[图片解析失败]"
                        else:
                            cell_img.paragraphs[0].text = "[无可用图片]"
                        cell_txt = tbl.cell(0, 1)
                        cell_txt.paragraphs[0].add_run(f"⚔️ {row_data['当前招式/动作']}\n").bold = True
                        p_details = cell_txt.add_paragraph()
                        p_details.add_run(f"前置招式: {row_data['父级招式']}\n")
                        p_details.add_run(f"招式类型: {row_data['动作类型']}\n")
                        p_details.add_run(f"触发指令: {row_data['触发按键']}\n")
                        p_details.add_run(f"手感帧数: 动画全长 {row_data['动画总帧数']}F | 判定活跃 {row_data['伤害判定帧']}\n")
                        p_details.add_run(f"战斗判定: {row_data.get('战斗判定/属性', '无')}\n")
                        p_details.add_run(f"美术交互: {row_data.get('设计备注', '无')}")
                        doc.add_paragraph().paragraph_format.space_after = Pt(10)
    return doc

def create_report_pdf(selected_configs):
    page_map = {}
    dummy_buf = io.BytesIO()
    build_pdf_story(dummy_buf, selected_configs, page_map)
    real_buf = io.BytesIO()
    build_pdf_story(real_buf, selected_configs, page_map)
    return real_buf.getvalue()

def build_pdf_story(buf, selected_configs, page_map):
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, PageBreak, Table, TableStyle, KeepTogether, Flowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas
    font_name = register_chinese_font()
    
    class RecordPage(Flowable):
        def __init__(self, key, p_map):
            Flowable.__init__(self)
            self.key = key
            self.p_map = p_map
        def draw(self):
            self.p_map[self.key] = self.canv.getPageNumber() - 1
        def wrap(self, availWidth, availHeight):
            return 0, 0
    
    class NumberedCanvas(canvas.Canvas):
        def __init__(self, *args, **kwargs):
            canvas.Canvas.__init__(self, *args, **kwargs)
        def showPage(self):
            self.draw_page_number()
            canvas.Canvas.showPage(self)
        def draw_page_number(self):
            self.saveState()
            if self._pageNumber > 1:
                try:
                    self.setFont(font_name, 9)
                except Exception:
                    self.setFont("Helvetica", 9)
                self.setFillColor(colors.HexColor('#7F8C8D'))
                display_page = self._pageNumber - 1
                self.drawCentredString(297.6, 20, f"第 {display_page} 页")
            self.restoreState()

        def setFont(self, psname, size, leading=None):
            canvas.Canvas.setFont(self, psname, size, leading)

    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36
    )
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        name='ExportBodyStyle',
        fontName=font_name,
        fontSize=10,
        leading=14,
        spaceAfter=5,
        textColor='#2C3E50'
    )
    table_body_text_style = ParagraphStyle(
        name='TableBodyText',
        parent=body_style,
        fontName=font_name,
        fontSize=8,
        leading=11,
        textColor='#2C3E50'
    )
    story = []
    tid_map = {
        "系统层级关系表": "system",
        "功能时序流向表": "seq",
        "有限状态机表": "fsm",
        "界面层级规划表": "ui",
        "战斗系统连招表": "combat",
        "鱼骨问题分析表": "fishbone"
    }
    story.append(Paragraph("<b>目 录</b>", ParagraphStyle(
        name='TOC_Title',
        fontName=font_name,
        fontSize=20,
        leading=24,
        alignment=1,
        spaceBefore=24,
        spaceAfter=18,
        textColor='#2C3E50'
    )))
    for item in selected_configs:
        raw_title = item["title"]
        parsed_title = raw_title[:-1] if raw_title.endswith("表") else raw_title
        current_tid = tid_map.get(raw_title, "system")
        scheme_title = st.session_state.get(f"scheme_title_{current_tid}", "").strip() or "游戏策划设计方案"
        main_key = f"{current_tid}_main"
        p_num = page_map.get(main_key, "1")
        story.append(Paragraph(
            f"<b>■ {scheme_title} ({parsed_title})</b><font color='#BDC3C7'> ............................................................................................................ </font><b> {p_num}</b>",
            ParagraphStyle(
                name=f'TOC_M_{parsed_title}_{p_num}',
                fontName=font_name,
                fontSize=11,
                leading=16,
                leftIndent=20,
                spaceBefore=8,
                spaceAfter=4,
                textColor='#2E4053'
            )
        ))
        sub_style = ParagraphStyle(
            name=f'TOC_Sub_{parsed_title}_{item["title"]}',
            fontName=font_name,
            fontSize=9.5,
            leading=14,
            leftIndent=40,
            spaceAfter=2,
            textColor='#7F8C8D'
        )
        
        sum_label = "问题总述" if current_tid == "fishbone" else "设计总述"
        sch_label = "迭代方案" if current_tid == "fishbone" else "设计方案"

        sub_sections = [
            (sum_label, f"{current_tid}_summary"),
            (sch_label, f"{current_tid}_scheme"),
            ("数据表单详细记录", f"{current_tid}_data")
        ]
        if item["graph_fn"]():
            sub_sections.append(("实时架构与关系流向预览", f"{current_tid}_graph"))
        for idx_sub, (sub_name, sub_key) in enumerate(sub_sections):
            sub_p_num = page_map.get(sub_key, "1")
            prefix = "├─ " if idx_sub < len(sub_sections) - 1 else "└─ "
            story.append(Paragraph(
                f"{prefix}{sub_name}<font color='#D5D8DC'> ................................................................................................................. </font><b> {sub_p_num}</b>",
                ParagraphStyle(
                    name=f'TOC_Sub_{parsed_title}_{sub_key}_{sub_p_num}',
                    parent=sub_style
                )
            ))
    story.append(PageBreak())
    for idx, item in enumerate(selected_configs):
        if idx > 0:
            story.append(PageBreak())
        raw_title = item["title"]
        parsed_title = raw_title[:-1] if raw_title.endswith("表") else raw_title
        current_tid = tid_map.get(raw_title, "system")
        scheme_title = st.session_state.get(f"scheme_title_{current_tid}", "").strip() or "游戏策划设计方案"
        story.append(RecordPage(f"{current_tid}_main", page_map))
        main_title_style = ParagraphStyle(
            name=f'ExportMainTitle_{idx}',
            fontName=font_name,
            fontSize=22,
            leading=28,
            alignment=1,
            spaceAfter=4,
            textColor='#2C3E50'
        )
        story.append(Paragraph(f"<b>{scheme_title}</b>", main_title_style))
        sub_title_style = ParagraphStyle(
            name=f'ExportSubTitle_{idx}',
            fontName=font_name,
            fontSize=13,
            leading=17,
            alignment=1,
            spaceAfter=12,
            textColor='#7F8C8D'
        )
        story.append(Paragraph(parsed_title, sub_title_style))
        
        owner_name = st.session_state.get(f"owner_{current_tid}", "").strip()
        owner_style = ParagraphStyle(
            name=f'ExportOwner_{idx}',
            fontName=font_name,
            fontSize=11,
            leading=15,
            alignment=2,
            spaceAfter=4,
            textColor='#7F8D8D'
        )
        story.append(Paragraph(f"<i>负责人：{owner_name if owner_name else '未指定'}</i>", owner_style))

        time_style = ParagraphStyle(
            name=f'ExportTime_{idx}',
            fontName=font_name,
            fontSize=11,
            leading=15,
            alignment=2,
            spaceAfter=18,
            textColor='#7F8D8D'
        )
        story.append(Paragraph(f"<i>迭代时间：{datetime.datetime.now().strftime('%Y/%m/%d')}</i>", time_style))

        h1_style = ParagraphStyle(
            name=f'ExportH1_{idx}',
            fontName=font_name,
            fontSize=14,
            leading=18,
            spaceBefore=12,
            spaceAfter=6,
            textColor='#2C3E50'
        )
        
        sum_label = "问题总述" if current_tid == "fishbone" else "设计总述"
        sch_label = "迭代方案" if current_tid == "fishbone" else "设计方案"

        story.append(RecordPage(f"{current_tid}_summary", page_map))
        story.append(Paragraph(f"<b>{sum_label}</b>", h1_style))
        
        ver = st.session_state.get(f"{current_tid}_ver", 0)
        summary_content = st.session_state.get(f"rte_{current_tid}_summary_text", "")
        if not summary_content or str(summary_content).strip() == "":
            summary_content = st.session_state.get(f"quill_{current_tid}_summary_{ver}", "")
        pdf_summary_html = md_to_pdf_html(summary_content)
        story.append(Paragraph(pdf_summary_html, body_style))
        story.append(Spacer(1, 10))
        story.append(RecordPage(f"{current_tid}_scheme", page_map))
        story.append(Paragraph(f"<b>{sch_label}</b>", h1_style))
        scheme_content = st.session_state.get(f"rte_{current_tid}_scheme_text", "")
        if not scheme_content or str(scheme_content).strip() == "":
            scheme_content = st.session_state.get(f"quill_{current_tid}_scheme_{ver}", "")
        pdf_scheme_html = md_to_pdf_html(scheme_content)
        story.append(Paragraph(pdf_scheme_html, body_style))
        story.append(Spacer(1, 15))
        story.append(RecordPage(f"{current_tid}_data", page_map))
        h2_data_banner_bytes, _, _ = cache_title_banner("📊 数据表单详细记录：", "H2")
        story.append(RLImage(io.BytesIO(h2_data_banner_bytes), width=480, height=30))
        story.append(Spacer(1, 6))
        df = clean_df_emojis(item["df"])
        table_data = []
        table_hdr_text_style = ParagraphStyle(
            name=f'TableHdrText_{idx}',
            parent=body_style,
            fontName=font_name,
            fontSize=9,
            leading=12,
            textColor='#FFFFFF',
            alignment=1
        )
        hdr_row = [Paragraph(f"<b>{col_name}</b>", table_hdr_text_style) for col_name in df.columns]
        table_data.append(hdr_row)
        for row in df.itertuples(index=False):
            row_cells = []
            for val in row:
                row_cells.append(Paragraph(str(val), table_body_text_style))
            table_data.append(row_cells)
        col_count = len(df.columns)
        col_widths = [523.0 / col_count] * col_count
        pdf_table = Table(table_data, colWidths=col_widths)
        pdf_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E4053')),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('LEFTPADDING', (0, 0), (-1, -1), 4),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#BDC3C7')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor('#F8F9F9'), colors.HexColor('#FFFFFF')])
        ]))
        story.append(pdf_table)
        graph_bytes = item["graph_fn"]()
        if graph_bytes:
            h2_graph_banner_bytes, _, _ = cache_title_banner("🎨 实时架构与关系流向预览：", "H2")
            rl_graph_title = RLImage(io.BytesIO(h2_graph_banner_bytes), width=520, height=33)
            image_stream = io.BytesIO(graph_bytes)
            try:
                pil_img = PILImage.open(image_stream)
                w, h = pil_img.size
                max_w = 520.0
                scale = min(1.0, max_w / float(w))
                if current_tid == "seq" and len(df) <= 6 and float(h) > 0:
                    scale = min(scale, 350.0 / float(h))
                new_w = w * scale
                new_h = h * scale
                image_stream.seek(0)
                rl_graph_img = RLImage(image_stream, width=new_w, height=new_h)
            except Exception:
                rl_graph_img = Paragraph("[未检测到本地 Graphviz 绘图环境，预览图已被省略]", body_style)
            story.append(Spacer(1, 15))
            if current_tid == "seq" and len(df) > 6:
                story.append(RecordPage(f"{current_tid}_graph", page_map))
                story.append(rl_graph_title)
                story.append(Spacer(1, 8))
                story.append(rl_graph_img)
            else:
                story.append(KeepTogether([
                    RecordPage(f"{current_tid}_graph", page_map),
                    rl_graph_title,
                    Spacer(1, 8),
                    rl_graph_img
                ]))
        if raw_title == "系统层级关系表":
            sun_bytes = compile_sunburst_image_bytes()
            if sun_bytes:
                h2_sun_banner_bytes, _, _ = cache_title_banner("🎯 架构权重分析 (Sunburst 占比图)：", "H2")
                rl_sun_title = RLImage(io.BytesIO(h2_sun_banner_bytes), width=520, height=29)
                rl_sun_img = RLImage(io.BytesIO(sun_bytes), width=300, height=300)
                story.append(Spacer(1, 20))
                story.append(KeepTogether([
                    rl_sun_title,
                    Spacer(1, 8),
                    rl_sun_img
                ]))
        elif raw_title == "界面层级规划表":
            valid_gallery = []
            for _, row_data in df.iterrows():
                ref_url = row_data.get("参考图URL")
                if pd.notna(ref_url) and str(ref_url).strip():
                    valid_gallery.append(row_data)
            if valid_gallery:
                h2_gal_banner_bytes, _, _ = cache_title_banner("🖼️ 界面卡片详细图鉴：", "H2")
                rl_gal_title = RLImage(io.BytesIO(h2_gal_banner_bytes), width=480, height=36)
                story.append(Spacer(1, 20))
                story.append(rl_gal_title)
                story.append(Spacer(1, 8))
                for row_data in valid_gallery:
                    img_path = get_local_image_path(str(row_data["参考图URL"]).strip())
                    if img_path and os.path.exists(img_path):
                        try:
                            rl_img = RLImage(img_path, width=160, height=120)
                        except Exception:
                            rl_img = Paragraph("[图片载入失败]", body_style)
                    else:
                        rl_img = Paragraph("[无图片]", body_style)
                    detail_text = f"""
                    <b>🏷️ {row_data['当前UI/界面']}</b><br/>
                    上级容器: {row_data['父级UI']}<br/>
                    界面类型: {row_data['界面类型']}<br/>
                    核心组件: {row_data.get('核心组件/元素', '无')}<br/>
                    设计备注: {row_data.get('界面备注', '无')}
                    """
                    detail_p = Paragraph(detail_text, table_body_text_style)
                    item_table = Table([[rl_img, detail_p]], colWidths=[180, 343])
                    item_table.setStyle(TableStyle([
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#BDC3C7')),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('TOPPADDING', (0, 0), (-1, -1), 6),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ]))
                    story.append(KeepTogether([
                        item_table,
                        Spacer(1, 10)
                    ]))
        elif raw_title == "战斗系统连招表":
            valid_gallery = []
            for _, row_data in df.iterrows():
                ref_url = row_data.get("招式参考图URL")
                if pd.notna(ref_url) and str(ref_url).strip():
                    valid_gallery.append(row_data)
            if valid_gallery:
                h2_gal_banner_bytes, _, _ = cache_title_banner("⚔️ 招式动作详细图鉴：", "H2")
                rl_gal_title = RLImage(io.BytesIO(h2_gal_banner_bytes), width=480, height=36)
                story.append(Spacer(1, 20))
                story.append(rl_gal_title)
                story.append(Spacer(1, 8))
                for row_data in valid_gallery:
                    img_path = get_local_image_path(str(row_data["招式参考图URL"]).strip())
                    if img_path and os.path.exists(img_path):
                        try:
                            rl_img = RLImage(img_path, width=160, height=120)
                        except Exception:
                            rl_img = Paragraph("[图片载入失败]", body_style)
                    else:
                        rl_img = Paragraph("[无图片]", body_style)
                    detail_text = f"""
                    <b>⚔️ {row_data['当前招式/动作']}</b> ({row_data['动作类型']})<br/>
                    前置招式: {row_data['父级招式']}<br/>
                    指令/按键: {row_data['触发按键']}<br/>
                    手感帧数: 动画全长 {row_data['动画总帧数']}F | 伤害判定 {row_data['伤害判定帧']}<br/>
                    战斗判定: {row_data.get('战斗判定/属性', '无')}<br/>
                    交互备注: {row_data.get('设计备注', '无')}
                    """
                    detail_p = Paragraph(detail_text, table_body_text_style)
                    item_table = Table([[rl_img, detail_p]], colWidths=[180, 343])
                    item_table.setStyle(TableStyle([
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#BDC3C7')),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('TOPPADDING', (0, 0), (-1, -1), 6),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ]))
                    story.append(KeepTogether([
                        item_table,
                        Spacer(1, 10)
                    ]))
    doc.build(story, canvasmaker=NumberedCanvas)

@st.dialog("📄 导出报告文档")
def show_export_dialog():
    st.write("请在下方定制您想要导出的文档属性：")
    export_format = st.radio("选择导出格式", ["Word 文档 (.docx)", "PDF 档案 (.pdf)"], on_change=reset_export_state)
    st.markdown("---")
    exp_sys = st.checkbox("🎮 系统层级关系表", value=True, on_change=reset_export_state)
    exp_seq = st.checkbox("⏱️ 功能时序流向表", value=False, on_change=reset_export_state)
    exp_fsm = st.checkbox("📋 有限状态机表", value=False, on_change=reset_export_state)
    exp_ui = st.checkbox("🎨 界面层级规划表", value=False, on_change=reset_export_state)
    exp_combat = st.checkbox("⚔️ 战斗系统连招表", value=False, on_change=reset_export_state)
    exp_fb = st.checkbox("🐟 鱼骨问题分析表", value=False, on_change=reset_export_state)
    st.markdown("---")
    selected_configs = []
    if exp_sys:
        selected_configs.append({"title": "系统层级关系表", "df": st.session_state.system_df, "graph_fn": compile_system_graph_bytes})
    if exp_seq:
        selected_configs.append({"title": "功能时序流向表", "df": st.session_state.seq_df, "graph_fn": compile_seq_graph_bytes})
    if exp_fsm:
        selected_configs.append({"title": "有限状态机表", "df": st.session_state.fsm_df, "graph_fn": compile_fsm_graph_bytes})
    if exp_ui:
        selected_configs.append({"title": "界面层级规划表", "df": st.session_state.ui_df, "graph_fn": compile_ui_graph_bytes})
    if exp_combat:
        selected_configs.append({"title": "战斗系统连招表", "df": st.session_state.combat_df, "graph_fn": compile_combat_graph_bytes})
    if exp_fb:
        selected_configs.append({"title": "鱼骨问题分析表", "df": st.session_state.fishbone_df, "graph_fn": compile_fishbone_graph_bytes})
    if not selected_configs:
        st.warning("⚠️ 请至少选择一个需要导出的系统记录工具。")
        reset_export_state()
        return
    main_container = st.empty()
    with main_container.container():
        if not st.session_state.export_ready:
            st.info("💡 提示：您已经修改了导出条件，请先点击下方“构建”按钮，待编译成功后再进行下载。")
        if st.button("⚙️ 开始构建报告", width="stretch"):
            with st.spinner("正在烘焙数据并执行多工具文档排版，请稍候..."):
                if export_format == "Word 文档 (.docx)":
                    try:
                        import docx
                        word_file_bytes = io.BytesIO()
                        doc_object = create_report_word(selected_configs)
                        doc_object.save(word_file_bytes)
                        st.session_state.export_bytes = word_file_bytes.getvalue()
                        st.session_state.export_filename = "Game_Designer_Report.docx"
                        st.session_state.export_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        st.session_state.export_ready = True
                    except ImportError:
                        st.error("❌ 无法导出 Word：您的执行环境中没有安装 `python-docx`。请执行 `pip install python-docx` 命令后再试。")
                        return
                else:
                    try:
                        from reportlab.platypus import SimpleDocTemplate
                        pdf_file_bytes = create_report_pdf(selected_configs)
                        st.session_state.export_bytes = pdf_file_bytes
                        st.session_state.export_filename = "Game_Designer_Report.pdf"
                        st.session_state.export_mime = "application/pdf"
                        st.session_state.export_ready = True
                    except ImportError:
                        st.error("❌ 无法导出 PDF：您的执行环境中没有安装 `reportlab`。请执行 `pip install reportlab` 命令后再试。")
                        return
        if st.session_state.export_ready:
            main_container.empty()
            with main_container.container():
                st.success("✅ 报告编译成功！已完美合并您勾选的所有模块和对应树图。")
                st.download_button(
                    label=f"📥 确认并下载 {export_format.split(' ')[0]}",
                    data=st.session_state.export_bytes,
                    file_name=st.session_state.export_filename,
                    mime=st.session_state.export_mime,
                    width="stretch"
                )
        if st.button("🔄 重新编译/更换选择", width="stretch"):
            reset_export_state()
            st.rerun()

# 辅助回调函数：将动作标记稳定、同步追加到全局和动态文本 Key 中，规避旧输入覆盖
def append_editor_text(key_text, textarea_key, value_to_append):
    current_text = st.session_state.get(textarea_key, "")
    new_text = current_text + value_to_append
    st.session_state[key_text] = new_text
    st.session_state[textarea_key] = new_text
    reset_export_state()

def draw_rich_text_editor(label_id, title_text):
    st.markdown(f"### {title_text}")
    key_text = f"rte_{label_id}_text"
    _tid = label_id.split("_")[0]
    ver = st.session_state.get(f"{_tid}_ver", 0)
    if HAS_QUILL:
        quill_key = f"quill_{label_id}_{ver}"
        if quill_key not in st.session_state:
            st.session_state[quill_key] = st.session_state[key_text]
        content = st_quill(
            value=st.session_state[quill_key],
            placeholder="在此编辑内容...可直接在下方更改字体大小、加粗、颜色、列表等效果，支持像 Word 一样实时渲染预览！",
            toolbar=[
                [{"size": []}],
                ["bold", "italic", "underline", "strike"],
                [{"color": []}, {"background": []}],
                [{"list": "ordered"}, {"list": "bullet"}],
                ["clean"]
            ],
            key=quill_key
        )
        just_imported_key = f"quill_just_imported_{label_id}"
        if content != st.session_state[key_text]:
            if st.session_state.get(just_imported_key, False):
                if content in (None, "", "<p><br></p>"):
                    pass
                else:
                    st.session_state[key_text] = content
                    st.session_state[just_imported_key] = False
                    reset_export_state()
            else:
                st.session_state[key_text] = content
                reset_export_state()
        if content and content not in (None, "", "<p><br></p>"):
            st.session_state[just_imported_key] = False
    else:
        st.markdown('<div class="editor-toolbar">', unsafe_allow_html=True)
        col_group = st.columns(10, gap="small")  # 升级为10列以平铺10个按键
        actions = [
            ("加粗 **B**", "**加粗文本**", f"b_{label_id}"),
            ("斜体 *I*", "*斜体文本*", f"i_{label_id}"),
            ("下划线 <u>", "<u>下划线文本</u>", f"u_{label_id}"),
            ("删除线 ~~", "~~删除线文本~~", f"s_{label_id}"),
            ("项目 •", "\n- 列表项目", f"bul_{label_id}"),
            ("编号 1.", "\n1. 列表项目", f"num_{label_id}"),
            ("字号 A+", "<font size='5'>18px大字号文本</font>", f"sz_{label_id}"),
            ("字色 🎨", "<font color='#E74C3C'>红字文本</font>", f"co_{label_id}"),  # 修复补齐字色按键
            ("雅黑 F", "<font face='Microsoft YaHei'>微软雅黑文字</font>", f"ft_{label_id}"),
            ("换行 ↩", "\n", f"br_{label_id}")
        ]
        textarea_key = f"{key_text}_{ver}"
        if textarea_key not in st.session_state:
            st.session_state[textarea_key] = st.session_state[key_text]
        
        for idx, (label, value_to_append, btn_key) in enumerate(actions):
            with col_group[idx]:
                st.button(
                    label, 
                    key=btn_key, 
                    on_click=append_editor_text, 
                    args=(key_text, textarea_key, value_to_append)
                )
        st.markdown('</div>', unsafe_allow_html=True)
        input_text = st.text_area(
            label="编辑区：",
            key=textarea_key,
            height=140,
            placeholder="请在此输入系统设计细节描述，支持 Markdown 与部分 HTML 标签混合排版...",
            label_visibility="collapsed"
        )
        if input_text != st.session_state[key_text]:
            st.session_state[key_text] = input_text
            reset_export_state()
        with st.expander("👁️ 查看实时排版解析预览：", expanded=True):
            if input_text:
                st.markdown(input_text, unsafe_allow_html=True)
            else:
                st.caption("暂无内容，请在编辑框内填写。")
        st.info("💡 **想要更棒的 Word 级编辑体验？** \n\n系统检测到未安装 `streamlit-quill` 模块。安装后，无需编写代码，**直接在输入框内框选文字，即可像使用 Word/WPS 一样点选字色、背景、字号并实时生效**！建议在您的终端命令行执行：\n\n```bash\npip install streamlit-quill\n```\n\n安装完成后重启本工具箱即可完美解锁。")

with st.sidebar:
    st.header("🛠️ 游戏设计工具箱")
    tool_mode = st.selectbox(
        "选择您需要使用的工具：",
        [
            "🎮 系统层级架构拆解",
            "⏱️ 功能时序流向拆解",
            "📋 有限状态机拆解",
            "🎨 界面层级拆解",
            "⚔️ 战斗系统连招拆解",
            "🐟 鱼骨问题分析"
        ]
    )
    st.markdown("---")
    st.subheader("👥 负责人管理")
    mode_to_tid = {
        "🎮 系统层级架构拆解": "system",
        "⏱️ 功能时序流向拆解": "seq",
        "📋 有限状态机拆解": "fsm",
        "🎨 界面层级拆解": "ui",
        "⚔️ 战斗系统连招拆解": "combat",
        "🐟 鱼骨问题分析": "fishbone"
    }
    active_tid = mode_to_tid[tool_mode]
    
    st.text_input(
        "📝 点击添加/修改负责人：",
        key=f"owner_{active_tid}",
        placeholder="例如：伍策划"
    )
    st.markdown("---")
    st.subheader("💾 本地数据管理")
    col_save, col_load = st.columns(2)
    with col_save:
        if st.button("💾 保存到本地", width="stretch"):
            show_save_dialog()
    with col_load:
        if st.button("📂 加载本地文档", width="stretch"):
            show_load_dialog()
    st.markdown("---")
    st.subheader("📄 导出报告管理")
    if st.button("📄 导出报告文档", width="stretch"):
        show_export_dialog()
    st.markdown("---")
    st.subheader("🔒 全局权限控制")
    is_edit_mode = st.toggle("✍️ 开启编辑模式", value=False)
    st.markdown("---")

if tool_mode == "🎮 系统层级架构拆解":
    st.title("🎮 游戏系统架构与层级拆解工具")
    st.markdown("本工具适用于拆分游戏系统层级关系，生成可视化的架构图。")
    
    st.text_input(
        "📝 方案标题",
        placeholder="请输入本设计方案的标题（如：英雄天赋系统设计方案）",
        key="scheme_title_system",
        on_change=reset_export_state
    )
    
    draw_rich_text_editor("system_summary", "设计总述")
    draw_rich_text_editor("system_scheme", "设计方案")
    st.markdown("---")
    with st.sidebar:
        st.subheader("🎨 节点类型与调色盘管理")
        if "node_types" not in st.session_state:
            st.session_state.node_types = pd.DataFrame([
                {"节点类型": "核心系统"},
                {"节点类型": "子系统"},
                {"节点类型": "功能"},
                {"节点类型": "玩法"},
                {"节点类型": "其他"}
            ])
        edited_types_df = st.data_editor(
            st.session_state.node_types,
            num_rows="dynamic" if is_edit_mode else "fixed",
            width="stretch",
            disabled=not is_edit_mode,
            column_config={"节点类型": st.column_config.TextColumn("类型名称", required=True)}
        )
        st.session_state.node_types = edited_types_df.dropna()
        dynamic_types_list = st.session_state.node_types["节点类型"].tolist()
        st.markdown("---")
        rankdir = st.selectbox("图表排版方向", options=["LR", "TB", "RL", "BT"], format_func=lambda x: {"LR": "从左到右 (横向)", "TB": "从上到下 (纵向)","RL": "从右到左 (横向)", "BT": "从下到上 (纵向)"}.get(x, x))
        st.markdown("##### ↔️ 界面板块自适应布局")
        ratio = st.slider("拖动滑块无缝调整 【数据编辑表】与【实时架构图】的显示比例：", min_value=20, max_value=80, value=45, step=5, key="ratio_system")
    tab1, tab2 = st.tabs(["📊 系统关系编辑与可视化", "🎯 架构权重分析 (旭日图)"])
    with tab1:
        col_data, col_viz = st.columns([ratio, 100 - ratio])
        with col_data:
            st.subheader("📝 系统层级关系表")
            edited_df = st.data_editor(
                st.session_state.system_df,
                key=f"system_editor_{st.session_state.system_ver}",
                num_rows="dynamic" if is_edit_mode else "fixed",
                width="stretch",
                disabled=not is_edit_mode,
                column_config={
                    "父系统": st.column_config.TextColumn("父节点/上级系统", required=True),
                    "子系统/模块": st.column_config.TextColumn("节点名称", required=True),
                    "节点类型": st.column_config.SelectboxColumn("节点类型", options=dynamic_types_list, required=True),
                    "复杂度权重": st.column_config.NumberColumn("复杂度权重 (1-5)", min_value=1, max_value=5, default=3),
                    "功能描述": st.column_config.TextColumn("功能简述")
                }
            )
            st.session_state.system_df = edited_df
            st.markdown("##### 📥 导出配置数据")
            export_data_section(edited_df, "system", "system_hierarchy", "系统架构")
        with col_viz:
            st.subheader("🌲 系统架构树状图")
            if edited_df.empty:
                st.warning("暂无数据。")
            else:
                png_preview = cache_system_graph(edited_df.to_json(orient="records"), rankdir)
                render_bytes_graph_actions(png_preview, "system", "system_hierarchy")
    with tab2:
        st.subheader("🎯 系统复杂度与权重占比分析")
        if not edited_df.empty:
            sun_bytes = compile_sunburst_image_bytes()
            if sun_bytes:
                st.image(sun_bytes, width="stretch")
            else:
                st.info("数据层级不足以生成占比图，请输入具有正确上下级关系的数据。")

elif tool_mode == "⏱️ 功能时序流向拆解":
    st.title("⏱️ 系统功能时序流向拆解工具")
    st.markdown("通过从上到下的 Excel 结构，快速直观地生成对应的系统功能流向顺序大纲。")
    
    st.text_input(
        "📝 方案标题",
        placeholder="请输入本时序流向设计方案的标题",
        key="scheme_title_seq",
        on_change=reset_export_state
    )
    
    draw_rich_text_editor("seq_summary", "设计总述")
    draw_rich_text_editor("seq_scheme", "设计方案")
    st.markdown("---")
    with st.sidebar:
        st.subheader("⚙️ 排版方向设置")
        seq_direction = st.selectbox("时序展示流向", options=["TB", "LR"], format_func=lambda x: "从上到下 (纵向时序)" if x == "TB" else "从左到右 (横向时序)")
        st.markdown("##### ↔️ 界面板块自适应布局")
        ratio = st.slider("拖动滑块调整比例：", min_value=20, max_value=80, value=45, key="ratio_seq")
    col_data, col_viz = st.columns([ratio, 100 - ratio])
    with col_data:
        st.subheader("📝 时序步骤表")
        seq_df = st.data_editor(
            st.session_state.seq_df,
            key=f"seq_editor_{st.session_state.seq_ver}",
            num_rows="dynamic" if is_edit_mode else "fixed",
            width="stretch",
            disabled=not is_edit_mode,
            column_config={
                "步骤": st.column_config.NumberColumn("时序号", required=True),
                "玩家输入": st.column_config.TextColumn("玩家操作/输入", required=True),
                "功能输出": st.column_config.TextColumn("系统响应/输出", required=True),
                "特殊表现": st.column_config.TextColumn("音效/特效/美术表现")
            }
        )
        st.session_state.seq_df = seq_df
        st.markdown("##### 📥 导出配置数据")
        export_data_section(seq_df, "seq", "sequence_data", "时序拆解")
    with col_viz:
        st.subheader("🕒 时序流图")
        if seq_df.empty:
            st.warning("暂无数据。")
        else:
            png_preview = cache_seq_graph(seq_df.to_json(orient="records"), seq_direction)
            render_bytes_graph_actions(png_preview, "seq", "sequence_flow")

elif tool_mode == "📋 有限状态机拆解":
    st.title("📋 有限状态机拆解工具")
    st.markdown("用于定义角色的 AI 行为逻辑，直观展示状态变迁、触发条件以及执行的 Action 动作。")
    
    st.text_input(
        "📝 方案标题",
        placeholder="请输入本有限状态机设计方案的标题",
        key="scheme_title_fsm",
        on_change=reset_export_state
    )
    
    draw_rich_text_editor("fsm_summary", "设计总述")
    draw_rich_text_editor("fsm_scheme", "设计方案")
    st.markdown("---")
    with st.sidebar:
        st.subheader("⚙️ 状态机视图排版")
        fsm_dir = st.selectbox("状态机排版方向", options=["LR", "TB"], format_func=lambda x: "从左到右 (横向流)" if x == "LR" else "从上到下 (纵向流)")
        st.markdown("##### ↔️ 界面板块自适应布局")
        ratio = st.slider("拖动滑块调整比例：", min_value=20, max_value=80, value=45, key="ratio_fsm")
    col_data, col_viz = st.columns([ratio, 100 - ratio])
    with col_data:
        st.subheader("📝 FSM 状态转换表")
        fsm_df = st.data_editor(
            st.session_state.fsm_df,
            key=f"fsm_editor_{st.session_state.fsm_ver}",
            num_rows="dynamic" if is_edit_mode else "fixed",
            width="stretch",
            disabled=not is_edit_mode,
            column_config={
                "源状态": st.column_config.TextColumn("当前状态 (Source State)", required=True),
                "源状态类型": st.column_config.SelectboxColumn("源状态类型", options=["初始状态", "常规状态", "结束状态"], default="常规状态"),
                "触发事件/条件": st.column_config.TextColumn("触发条件 (Transition Event)", required=True),
                "变迁动作 (Action)": st.column_config.TextColumn("执行动作 (Action)"),
                "目标状态": st.column_config.TextColumn("变迁目标 (Target State)", required=True),
                "目标状态类型": st.column_config.SelectboxColumn("目标状态类型", options=["初始状态", "常规状态", "结束状态"], default="常规状态")
            }
        )
        st.session_state.fsm_df = fsm_df
        st.markdown("##### 📥 导出配置数据")
        export_data_section(fsm_df, "fsm", "fsm_transitions", "FSM拆解")
    with col_viz:
        st.subheader("📋 FSM 状态转换流向图")
        if fsm_df.empty:
            st.warning("暂无数据。")
        else:
            png_preview = cache_fsm_graph(fsm_df.to_json(orient="records"), fsm_dir)
            render_bytes_graph_actions(png_preview, "fsm", "fsm_diagram")

elif tool_mode == "🎨 界面层级拆解":
    st.title("🎨 界面层级拆解工具")
    st.markdown("规划整个游戏的 UI 层级树状图，并可在右侧实时渲染并预览对应的美术/策划参考图。")
    
    st.text_input(
        "📝 方案标题",
        placeholder="请输入本界面层级规划设计方案的标题",
        key="scheme_title_ui",
        on_change=reset_export_state
    )
    
    draw_rich_text_editor("ui_summary", "设计总述")
    draw_rich_text_editor("ui_scheme", "设计方案")
    st.markdown("---")
    with st.sidebar:
        st.subheader("⚙️ 界面层级方向")
        ui_rankdir = st.selectbox("层级树排版方向", options=["LR", "TB"], format_func=lambda x: "从左到右 (横向树)" if x == "LR" else "从上到下 (纵向树)", key="ui_rankdir")
        st.markdown("##### ↔️ 界面板块自适应布局")
        ratio = st.slider("拖动滑块调整比例：", min_value=20, max_value=80, value=45, key="ratio_ui")
    col_data, col_viz = st.columns([ratio, 100 - ratio])
    with col_data:
        st.subheader("📝 界面层级规划表")
        ui_df = st.data_editor(
            st.session_state.ui_df,
            key=f"ui_editor_{st.session_state.ui_ver}",
            num_rows="dynamic" if is_edit_mode else "fixed",
            width="stretch",
            disabled=not is_edit_mode,
            column_config={
                "父级UI": st.column_config.TextColumn("父级容器/父界面", required=True),
                "当前UI/界面": st.column_config.TextColumn("界面/组件名称", required=True),
                "界面类型": st.column_config.SelectboxColumn("类型", options=["全屏主界面", "独立功能页", "二级弹窗", "浮窗/提示"], default="独立功能页"),
                "参考图URL": st.column_config.TextColumn("参考图 (本地绝对路径或网络URL)", help="直接渲染至图表内部"),
                "核心组件/元素": st.column_config.TextColumn("重要组件元素列表"),
                "界面备注": st.column_config.TextColumn("设计要点/交互规范")
            }
        )
        st.session_state.ui_df = ui_df
        st.markdown("##### 📥 导出配置数据")
        export_data_section(ui_df, "ui", "ui_hierarchy", "UI层级结构")
    with col_viz:
        st.subheader("🌲 UI树状拆解图")
        if ui_df.empty:
            st.warning("暂无数据。")
        else:
            tab_graph, tab_gallery = st.tabs(["🌲 界面层级架构图", "🖼️ 界面卡片详细图鉴"])
            with tab_graph:
                png_preview = cache_ui_graph(ui_df.to_json(orient="records"), ui_rankdir)
                render_bytes_graph_actions(png_preview, "ui", "ui_hierarchy")
            with tab_gallery:
                st.caption("📱 UI 参考大图及备注图鉴：")
                for _, row_data in ui_df.iterrows():
                    child_name = row_data["当前UI/界面"]
                    parent_name = row_data["父级UI"]
                    img_url = row_data["参考图URL"]
                    notes = row_data["界面备注"]
                    comps = row_data["核心组件/元素"]
                    if pd.notna(img_url) and str(img_url).strip():
                        with st.container():
                            col_img, col_txt = st.columns([1, 2])
                            with col_img:
                                st.image(str(img_url).strip(), caption=child_name, width="stretch")
                            with col_txt:
                                st.subheader(f"🏷️ {child_name}")
                                st.markdown(f"- **上级容器**: `{parent_name}`")
                                st.markdown(f"- **核心组件**: {comps if pd.notna(comps) else '无'}")
                                st.markdown(f"- **设计备注**: {notes if pd.notna(notes) else '无'}")
                                st.divider()

elif tool_mode == "⚔️ 战斗系统连招拆解":
    st.title("⚔️ 战斗系统动作派生与连招拆解工具")
    st.markdown("梳理战斗角色的招式判定、派生动作链条与手感帧，支持动作参考原画直观嵌入动作流中。")
    
    st.text_input(
        "📝 方案标题",
        placeholder="请输入本连招与动作派生设计方案的标题",
        key="scheme_title_combat",
        on_change=reset_export_state
    )
    
    draw_rich_text_editor("combat_summary", "设计总述")
    draw_rich_text_editor("combat_scheme", "设计方案")
    st.markdown("---")
    with st.sidebar:
        st.subheader("⚙️ 招式流向方向")
        combat_rankdir = st.selectbox("动作树排版方向", options=["LR", "TB"], format_func=lambda x: "从左到右 (横向树)" if x == "LR" else "从上到下 (纵向树)", key="combat_rankdir")
        st.markdown("##### ↔️ 界面板块自适应布局")
        ratio = st.slider("拖动滑块调整比例：", min_value=20, max_value=80, value=45, key="ratio_combat")
    col_data, col_viz = st.columns([ratio, 100 - ratio])
    with col_data:
        st.subheader("📝 招式派生表")
        combat_df = st.data_editor(
            st.session_state.combat_df,
            key=f"combat_editor_{st.session_state.combat_ver}",
            num_rows="dynamic" if is_edit_mode else "fixed",
            width="stretch",
            disabled=not is_edit_mode,
            column_config={
                "父级招式": st.column_config.TextColumn("起手/前置招式", required=True, help="例如: 普通攻击一段"),
                "当前招式/动作": st.column_config.TextColumn("派生/当前招式", required=True, help="例如: 普通攻击二段"),
                "动作类型": st.column_config.SelectboxColumn("动作类型", options=["起手招式", "派生连招", "终结技", "蓄力/特殊", "闪避/防御"], default="派生连招"),
                "触发按键": st.column_config.TextColumn("输入指令/条件", required=True, help="例如: J 键, 双击 J, 或 判定框碰撞"),
                "动画总帧数": st.column_config.NumberColumn("动画全长(F)", min_value=1, default=30),
                "伤害判定帧": st.column_config.TextColumn("伤害活性帧区间", help="例如: 10-15 或 多段: 8-10|15-18"),
                "战斗判定/属性": st.column_config.TextColumn("核心战斗机制", help="例如: 霸体 / 击飞 / 空中追击 / 削韧值"),
                "招式参考图URL": st.column_config.TextColumn("招式参考图 (本地路径或网络URL)", help="直接渲染到连招树中"),
                "设计备注": st.column_config.TextColumn("美术&音效反馈及交互说明")
            }
        )
        st.session_state.combat_df = combat_df
        st.markdown("##### 📥 导出配置数据")
        export_data_section(combat_df, "combat", "combat_hierarchy", "连招派生结构")
    with col_viz:
        st.subheader("🌲 动作连招派生树")
        if combat_df.empty:
            st.warning("暂无数据。")
        else:
            tab_graph, tab_gallery = st.tabs(["🌲 派生连招树状图", "🖼️ 招式动作详细图鉴"])
            with tab_graph:
                png_preview = cache_combat_graph(combat_df.to_json(orient="records"), combat_rankdir)
                render_bytes_graph_actions(png_preview, "combat", "combat_hierarchy")
            with tab_gallery:
                st.caption("⚔️ 派生树中动作美术参考大图及策划备注：")
                for _, row_data in combat_df.iterrows():
                    child_name = row_data["当前招式/动作"]
                    parent_name = row_data["父级招式"]
                    action_type = row_data["动作类型"]
                    key_input = row_data["触发按键"]
                    total_frames = row_data["动画总帧数"]
                    hit_frames = row_data["伤害判定帧"]
                    props = row_data["战斗判定/属性"]
                    img_url = row_data["招式参考图URL"]
                    notes = row_data["设计备注"]
                    if pd.notna(img_url) and str(img_url).strip():
                        with st.container():
                            col_img, col_txt = st.columns([1, 2])
                            with col_img:
                                st.image(str(img_url).strip(), caption=child_name, width="stretch")
                            with col_txt:
                                st.subheader(f"⚔️ {child_name}")
                                st.markdown(f"- **前置招式**: `{parent_name}`")
                                st.markdown(f"- **指令/按键**: `{key_input}`")
                                st.markdown(f"- **动画总帧数**: `{total_frames} F` | **伤害判定活跃帧**: `{hit_frames}`")
                                st.markdown(f"- **招式类型**: `{action_type}`")
                                st.markdown(f"- **战斗判定/机制属性**: `{props}`")
                                st.markdown(f"- **交互美术备注**: {notes if pd.notna(notes) else '无'}")
                                st.divider()

elif tool_mode == "🐟 鱼骨问题分析":
    st.title("🐟 鱼骨问题与因果拆解分析工具")
    st.markdown("本工具采用石川图/鱼骨图方法，用于多维度推导游戏研发与运营中遇到的问题根源并导出改进迭代方案。")
    
    st.text_input(
        "📝 问题核心/分析标题",
        placeholder="请输入需要分析的鱼骨图核心问题（例如：版本首发次日留存暴跌原因分析）",
        key="scheme_title_fishbone",
        on_change=reset_export_state
    )
    
    draw_rich_text_editor("fishbone_summary", "问题总述")
    draw_rich_text_editor("fishbone_scheme", "迭代方案")
    st.markdown("---")
    with st.sidebar:
        st.subheader("⚙️ 鱼骨图配置")
        st.markdown("##### ↔️ 界面板块自适应布局")
        ratio = st.slider("拖动滑块调整比例：", min_value=20, max_value=80, value=45, key="ratio_fishbone")
    col_data, col_viz = st.columns([ratio, 100 - ratio])
    with col_data:
        st.subheader("📝 鱼骨因果分析表")
        fishbone_df = st.data_editor(
            st.session_state.fishbone_df,
            key=f"fishbone_editor_{st.session_state.fishbone_ver}",
            num_rows="dynamic" if is_edit_mode else "fixed",
            width="stretch",
            disabled=not is_edit_mode,
            column_config={
                "大类/主骨": st.column_config.TextColumn("分析大类 (主骨)", required=True, help="例如: 客户端, 服务端, 策划设计, 运营"),
                "具体原因/小骨": st.column_config.TextColumn("具体子因 (小骨)", required=True, help="造成该大类问题的具体小原因"),
                "严重程度": st.column_config.NumberColumn("严重权重 (1-5)", min_value=1, max_value=5, default=3),
                "原因细述": st.column_config.TextColumn("原因详尽叙述与临时补救方案")
            }
        )
        st.session_state.fishbone_df = fishbone_df
        st.markdown("##### 📥 导出配置数据")
        export_data_section(fishbone_df, "fishbone", "fishbone_data", "问题根源分析")
    with col_viz:
        st.subheader("🐟 实时因果分析鱼骨图")
        if fishbone_df.empty:
            st.warning("暂无分析数据，请添加相应维度。")
        else:
            png_preview = cache_fishbone_graph(fishbone_df.to_json(orient="records"), st.session_state.get("scheme_title_fishbone", "").strip() or "未命名核心问题")
            render_bytes_graph_actions(png_preview, "fishbone", "fishbone_chart")